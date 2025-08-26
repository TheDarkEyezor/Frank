#!/usr/bin/env python3
"""
Check Resume Content
Examine the actual content of resume files to understand their structure.
"""

import os
from docx import Document

def check_resume_content():
    """Check the content of resume files"""
    print("📋 Checking Resume Content")
    print("=" * 50)
    
    resume_files = [
        "AdiPrabs_SWE.docx",
        "AdiPrabs_Quant.docx", 
        "AdiPrabs_Cons.docx"
    ]
    
    for resume_file in resume_files:
        print(f"\n📄 {resume_file}:")
        
        if not os.path.exists(resume_file):
            print(f"   ❌ File not found")
            continue
        
        try:
            doc = Document(resume_file)
            
            # Get text from paragraphs
            print(f"   📝 Paragraphs ({len(doc.paragraphs)}):")
            for i, paragraph in enumerate(doc.paragraphs[:5], 1):  # Show first 5
                text = paragraph.text.strip()
                if text:
                    print(f"      {i}. {text[:100]}...")
            
            # Get text from tables
            print(f"   📊 Tables ({len(doc.tables)}):")
            for i, table in enumerate(doc.tables[:3], 1):  # Show first 3 tables
                print(f"      Table {i}:")
                for j, row in enumerate(table.rows[:3], 1):  # Show first 3 rows
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text[:50])
                    if row_text:
                        print(f"         Row {j}: {' | '.join(row_text)}")
            
            # Check for placeholders
            all_text = ""
            for paragraph in doc.paragraphs:
                all_text += paragraph.text + " "
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + " "
            
            if "<<Program>>" in all_text:
                print(f"   ✅ Found <<Program>> placeholder")
            else:
                print(f"   ❌ No <<Program>> placeholder found")
                
            if "<<Company>>" in all_text:
                print(f"   ✅ Found <<Company>> placeholder")
            else:
                print(f"   ❌ No <<Company>> placeholder found")
                
        except Exception as e:
            print(f"   ❌ Error reading file: {e}")

if __name__ == "__main__":
    check_resume_content()

