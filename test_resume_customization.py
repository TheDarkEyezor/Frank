#!/usr/bin/env python3
"""
Test Resume Customization
Verify that text replacement works correctly in resume files.
"""

import os
import sys
from pathlib import Path

def test_resume_customization():
    """Test the resume customization functionality"""
    print("🧪 Testing Resume Customization")
    print("=" * 50)
    
    # Test cases
    test_cases = [
        {
            "resume_path": "AdiPrabs_SWE.docx",
            "company_name": "Test Company",
            "job_title": "Software Engineer"
        },
        {
            "resume_path": "AdiPrabs_Quant.docx", 
            "company_name": "Hedge Fund",
            "job_title": "Quantitative Analyst"
        },
        {
            "resume_path": "AdiPrabs_Cons.docx",
            "company_name": "Consulting Firm",
            "job_title": "Business Analyst"
        }
    ]
    
    results = []
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\n📋 Test {i}: {test_case['resume_path']}")
        print(f"   Company: {test_case['company_name']}")
        print(f"   Job Title: {test_case['job_title']}")
        
        # Check if file exists
        if not os.path.exists(test_case['resume_path']):
            print(f"   ❌ File not found: {test_case['resume_path']}")
            results.append(False)
            continue
        
        # Test customization
        try:
            customized_path = customize_resume_with_placeholders(
                test_case['resume_path'],
                test_case['company_name'],
                test_case['job_title']
            )
            
            if customized_path and os.path.exists(customized_path):
                print(f"   ✅ Customization successful: {customized_path}")
                
                # Verify the customization worked
                if verify_customization(customized_path, test_case['company_name'], test_case['job_title']):
                    print(f"   ✅ Verification passed")
                    results.append(True)
                else:
                    print(f"   ❌ Verification failed")
                    results.append(False)
            else:
                print(f"   ❌ Customization failed")
                results.append(False)
                
        except Exception as e:
            print(f"   ❌ Error during customization: {e}")
            results.append(False)
    
    # Summary
    print(f"\n📊 Test Results:")
    passed = sum(results)
    total = len(results)
    
    for i, result in enumerate(results):
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"   Test {i+1}: {status}")
    
    print(f"\n🎯 Overall: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All resume customization tests passed!")
        return True
    else:
        print("⚠️ Some tests failed. Resume customization may not work correctly.")
        return False

def customize_resume_with_placeholders(resume_path, company_name, job_title):
    """Customize resume by replacing <<Program>> and <<Company>> placeholders"""
    try:
        if not os.path.exists(resume_path):
            print(f"❌ Resume file not found: {resume_path}")
            return resume_path
        
        if not company_name and not job_title:
            print("⚠️ No company name or job title available for customization")
            return resume_path
        
        # Create customized filename
        base_name = os.path.splitext(resume_path)[0]
        extension = os.path.splitext(resume_path)[1]
        customized_path = f"{base_name}_{company_name.replace(' ', '')}{extension}"
        
        # For DOCX files, we need python-docx library
        try:
            from docx import Document
            
            # Load the document
            doc = Document(resume_path)
            
            replacements_made = 0
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                if "<<Program>>" in paragraph.text or "<<Company>>" in paragraph.text:
                    print(f"   📝 Found placeholders in paragraph: {paragraph.text[:100]}...")
                    
                    # Replace placeholders
                    new_text = paragraph.text
                    if job_title:
                        new_text = new_text.replace("<<Program>>", job_title)
                    if company_name:
                        new_text = new_text.replace("<<Company>>", company_name)
                    
                    # Clear and set new text
                    paragraph.clear()
                    paragraph.add_run(new_text)
                    replacements_made += 1
                    print(f"   ✅ Updated to: {new_text[:100]}...")
            
            # Replace placeholders in tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "<<Program>>" in cell.text or "<<Company>>" in cell.text:
                            print(f"   📝 Found placeholders in table cell: {cell.text[:50]}...")
                            
                            new_text = cell.text
                            if job_title:
                                new_text = new_text.replace("<<Program>>", job_title)
                            if company_name:
                                new_text = new_text.replace("<<Company>>", company_name)
                            
                            cell.text = new_text
                            replacements_made += 1
                            print(f"   ✅ Updated cell to: {new_text[:50]}...")
            
            if replacements_made > 0:
                # Save customized resume
                doc.save(customized_path)
                print(f"   💾 Customized resume saved as: {customized_path}")
                return customized_path
            else:
                print(f"   ⚠️ No placeholders found in resume")
                return resume_path
                
        except ImportError:
            print("⚠️ python-docx not installed. Install with: pip install python-docx")
            print("⚠️ Using original resume without customization")
            return resume_path
        except Exception as e:
            print(f"❌ Error customizing resume: {e}")
            return resume_path
            
    except Exception as e:
        print(f"❌ Error in resume customization: {e}")
        return resume_path

def verify_customization(customized_path, company_name, job_title):
    """Verify that the customization worked correctly"""
    try:
        from docx import Document
        
        doc = Document(customized_path)
        
        # Check if placeholders were replaced
        content = ""
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            content += paragraph.text + " "
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
        
        # Check if placeholders still exist
        if "<<Program>>" in content or "<<Company>>" in content:
            print(f"   ⚠️ Placeholders still found in content")
            return False
        
        # Check if replacements were made
        if job_title and job_title not in content:
            print(f"   ⚠️ Job title '{job_title}' not found in content")
            return False
        
        if company_name and company_name not in content:
            print(f"   ⚠️ Company name '{company_name}' not found in content")
            return False
        
        return True
        
    except Exception as e:
        print(f"   ❌ Error verifying customization: {e}")
        return False

if __name__ == "__main__":
    success = test_resume_customization()
    
    if success:
        print("\n✅ Resume customization is working correctly!")
        print("🚀 Proceeding with application filling...")
    else:
        print("\n❌ Resume customization has issues!")
        print("⚠️ Please fix the issues before proceeding.")
        sys.exit(1)

