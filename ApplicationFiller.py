responses = {
  "first name": "Aditya",
  "last name": "Prabakaran", 
  "full name": "Aditya Prabakaran",
  "email": "aditya.prabakaran@gmail.com",
  "phone": "+447587460771",
  "preferred contact method": "email",
  "preferred name": "Adi",
  "school": "Imperial College London",
  "university": "Imperial College London",
  "degree": "MEng Computer Science",
  "graduation year": "2024",
  "gpa": "First Class Honours",
  "city": "London",
  "country": "United Kingdom",
  "major": "Computer Science",
  "start date": "2020-09-01",
  "end date": "2024-05-15",
  "linkedin": "https://www.linkedin.com/in/adiprabs",
  "github": "https://github.com/TheDarkEyezor",
  "profile": "https://adiprabs.vercel.app/",
  "united states need sponsorship": "yes",
  "united kingdom need sponsorship": "no",
  "military service": "none",
  "disability": "none",
  "gender": "male",
  "sexuality": "straight/cis",
  
  # Work experience details
  "current company": "Trajex",
  "current position": "Machine Learning Developer",
  "current start date": "2024-10-01",
  "previous company": "Altus Reach", 
  "previous position": "Machine Learning Engineer",
  "previous start date": "2023-12-01",
  "previous end date": "2024-10-01",
  
  # Account creation credentials
  "password": "!n+3rn10|",
}

import os
import json
import requests
from pathlib import Path
import time

class ApplicationFiller:
  """
  Enhanced automated job application form filler with intelligent dropdown handling,
  Ollama LLM integration for RAG-based responses, and automatic file upload capabilities.
  
  Features:
  - Smart dropdown option matching with proper selection
  - Ollama LLM integration for unknown questions and customized responses
  - RAG document support for context-aware responses
  - Automatic resume/CV file upload
  - Visual form filling with real-time feedback (default: live mode)
  - Customized resume generation for specific job roles
  """
  
  def __init__(self, link="", model="llama3.2", headless=False, slow_mode=True, 
               ollama_base_url="http://localhost:11434", resume_path="sample_resume.txt", 
               reference_doc_path=None, company_name="", job_title=""):
    self.link = link
    self.model = model
    self.driver = None
    self.headless = headless  # Default to False for live viewing
    self.slow_mode = slow_mode  # Default to True for better visualization
    self.ollama_base_url = ollama_base_url
    self.resume_path = resume_path
    self.reference_doc_path = reference_doc_path
    self.company_name = company_name
    self.job_title = job_title
    self.reference_content = self._load_reference_document()
    self.resume_content = self._load_resume()
    
    # Updated resume paths for different job types (DOCX format)
    self.resume_paths = {
      "swe": "AdiPrabs_SWE.docx",
      "quant": "AdiPrabs_Quant.docx", 
      "communication": "AdiPrabs_Cons.docx"
    }
    
    # Verify resume files exist
    for job_type, path in self.resume_paths.items():
      if not os.path.exists(path):
        print(f"⚠️ Resume file not found: {path}")
    
    # Load profile for enhanced LLM responses
    self.profile_content = self._load_profile_document()
    
    # Website-specific configurations
    self.website_configs = {
      "greenhouse.io": {
        "type": "greenhouse",
        "requires_account": False,
        "custom_dropdowns": True,
        "enter_key_for_dropdowns": True,
        # Defer file uploads (resume) until after other fields to avoid interrupting
        # interactive widgets like React-select comboboxes
        "defer_file_uploads": True,
        "has_cookies": False,
        "extract_job_info_early": True
      },
      "temasek.com.sg": {
        "type": "temasek_careers",
        "requires_account": True,
        "apply_button_required": True,
        "account_creation_url": "/talentcommunity/apply/",
        "login_url": "/talentcommunity/login/",
        "email_field": "email",
        "password_field": "password",
        "has_cookies": True,
        "cookie_button_text": "Accept",
        "extract_job_info_early": True,
        "education_fields": True,
        "experience_fields": True
      },
      "squarepoint-capital.com": {
        "type": "squarepoint",
        "requires_account": True,
        "apply_button_required": True,
        "account_creation_url": "/careers/apply/",
        "login_url": "/careers/login/",
        "has_cookies": True,
        "cookie_button_text": "I agree to all cookie use",
        "extract_job_info_early": True,
        "education_fields": True,
        "experience_fields": True
      },
      "verition.com": {
        "type": "verition",
        "requires_account": True,
        "apply_button_required": True,
        "account_creation_url": "/careers/apply/",
        "login_url": "/careers/login/",
        "has_cookies": True,
        "extract_job_info_early": True,
        "education_fields": True,
        "experience_fields": True
      },
      "helsing.ai": {
        "type": "direct_form",
        "requires_account": False, 
        "apply_button_required": False,
        "form_selector": "form",
        "has_cookies": False,
        "extract_job_info_early": True
      },
      "citadel.com": {
        "type": "career_site",
        "requires_account": True,
        "apply_button_required": True,
        "account_creation_url": "/careers/apply/",
        "email_field": "email",
        "has_cookies": True,
        "cookie_button_text": "Accept All Cookies",
        "extract_job_info_early": True
      },
      "davincitrading.com": {
        "type": "career_site",
        "requires_account": True,
        "apply_button_required": True,
        "account_creation_url": "/careers/apply/",
        "email_field": "email",
        "has_cookies": True,
        "extract_job_info_early": True
      },
      "revolut.com": {
        "type": "career_portal",
        "requires_account": False,  # Direct application form
        "apply_button_required": False,  # Already on application page
        "apply_button_selector": "button:contains('Submit application')",
        "has_cookies": True,
        "cookie_button_text": "Accept All Cookies",
        "extract_job_info_early": True,
        "job_title_selector": "h1",
        "form_sections": ["General information", "Show us your work", "Experience", "Education"],
        "file_upload_selector": "input[type='file']",
        "dropdown_selectors": ["select", "div[role='combobox']", "button[role='combobox']"],
        "specific_fields": {
          "pronouns": "He/him",
          "formula1_experience": "No", 
          "previous_revolut_employee": "No",
          "interview_transcript_consent": "Yes, I consent"
        }
      }
    }    # Extract website type from URL
    self.current_website_config = self._detect_website_type()
    
  def handle_cookies(self):
    """Handle cookie consent banners on websites"""
    try:
      from selenium.webdriver.common.by import By
      from selenium.webdriver.support.ui import WebDriverWait
      from selenium.webdriver.support import expected_conditions as EC
      
      if not self.driver:
        return True
      
      print("🍪 Checking for cookie consent banners...")
      
      # Common cookie consent selectors
      cookie_selectors = [
        # Generic selectors
        "button:contains('Accept All')",
        "button:contains('Accept all')", 
        "button:contains('Accept All Cookies')",
        "button:contains('Accept')",
        "a:contains('Accept All')",
        
        # Specific selectors from websites
        "[data-testid*='accept']",
        "[data-qa*='accept']",
        ".accept-all",
        "#accept-all",
        ".cookie-accept",
        "#cookie-accept",
        
        # Revolut specific (from the fetched content)
        "button:contains('Accept All Cookies')",
        
        # Citadel specific (from the fetched content) 
        "button:contains('Accept All Cookies')",
        
        # Common patterns
        "[aria-label*='Accept']",
        "[title*='Accept']"
      ]
      
      cookie_button = None
      
      # Try to find cookie consent button
      for selector in cookie_selectors:
        try:
          if ":contains" in selector:
            # Use XPath for text-based selectors
            text_options = ["Accept All", "Accept all", "Accept All Cookies", "Accept"]
            for text in text_options:
              try:
                xpath_selector = f"//*[contains(text(), '{text}') and (self::button or self::a)]"
                cookie_button = self.driver.find_element(By.XPATH, xpath_selector)
                if cookie_button and cookie_button.is_displayed() and cookie_button.is_enabled():
                  break
              except:
                continue
          else:
            cookie_button = self.driver.find_element(By.CSS_SELECTOR, selector)
          
          if cookie_button and cookie_button.is_displayed() and cookie_button.is_enabled():
            break
        except:
          continue
      
      if cookie_button:
        print("🍪 Found cookie consent button - accepting...")
        if not self.headless:
          self.driver.execute_script("arguments[0].style.border='3px solid orange'", cookie_button)
          time.sleep(1)
        
        try:
          cookie_button.click()
          print("✅ Cookie consent accepted")
          time.sleep(2)  # Wait for banner to disappear
          return True
        except Exception as e:
          print(f"⚠️ Could not click cookie button: {e}")
          try:
            # Try JavaScript click
            self.driver.execute_script("arguments[0].click();", cookie_button)
            print("✅ Cookie consent accepted via JavaScript")
            time.sleep(2)
            return True
          except Exception as e2:
            print(f"⚠️ JavaScript click also failed: {e2}")
      else:
        print("🍪 No cookie consent banner found")
      
      return True
      
    except Exception as e:
      print(f"⚠️ Error handling cookies: {e}")
      return True  # Continue anyway
      
  def handle_redirects(self):
    """Handle unexpected redirects and ensure we're on the right page"""
    try:
      if not self.driver:
        return True
      
      current_url = self.driver.current_url
      original_domain = self.link.split('/')[2] if '/' in self.link else self.link
      current_domain = current_url.split('/')[2] if '/' in current_url else current_url
      
      # Check if we were redirected to a different domain
      if original_domain.lower() != current_domain.lower():
        print(f"⚠️ Redirect detected: {original_domain} → {current_domain}")
        
        # Check if it's a legitimate redirect within the same organization
        legitimate_redirects = {
          "revolut.com": ["app.revolut.com", "jobs.revolut.com", "careers.revolut.com"],
          "citadel.com": ["careers.citadel.com", "jobs.citadel.com"],
          "helsing.ai": ["jobs.helsing.ai", "careers.helsing.ai"],
          "davincitrading.com": ["careers.davincitrading.com", "jobs.davincitrading.com"]
        }
        
        is_legitimate = False
        for main_domain, allowed_redirects in legitimate_redirects.items():
          if main_domain in original_domain.lower():
            if any(allowed in current_domain.lower() for allowed in allowed_redirects):
              is_legitimate = True
              print(f"✅ Legitimate redirect within {main_domain} organization")
              break
        
        if not is_legitimate:
          print(f"⚠️ Unexpected redirect - may need manual intervention")
          
      # Update the link to current URL for consistency
      self.link = current_url
      return True
      
    except Exception as e:
      print(f"⚠️ Error handling redirects: {e}")
      return True
      
  def extract_job_information_early(self):
    """Extract job information before clicking apply buttons - this runs early"""
    try:
      from selenium.webdriver.common.by import By
      
      if not self.driver:
        return ""
      
      print("📋 Extracting job information from page...")
      
      # Look for job title and description on the current page
      job_info = ""
      
      # Try to extract job title
      title_selectors = [
        "h1",
        ".job-title", 
        "#job-title",
        "[data-qa*='title']",
        ".title",
        "h2"
      ]
      
      job_title = ""
      for selector in title_selectors:
        try:
          title_element = self.driver.find_element(By.CSS_SELECTOR, selector)
          if title_element and title_element.text.strip():
            job_title = title_element.text.strip()
            print(f"📝 Found job title: {job_title}")
            break
        except:
          continue
      
      # Update our stored job title if we found one
      if job_title and not self.job_title:
        self.job_title = job_title
      
      # Try to extract job description
      desc_selectors = [
        ".job-description",
        "#job-description", 
        "[data-qa*='description']",
        ".description",
        ".content",
        "main p",
        ".job-detail",
        ".job-content"
      ]
      
      job_description = ""
      for selector in desc_selectors:
        try:
          desc_elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
          for element in desc_elements:
            if element.text and len(element.text.strip()) > 100:
              job_description = element.text.strip()[:2000]  # Limit length
              print(f"📝 Found job description: {len(job_description)} characters")
              break
          if job_description:
            break
        except:
          continue
      
      # If no specific description found, try to get page content
      if not job_description:
        try:
          body_text = self.driver.find_element(By.TAG_NAME, "body").text
          # Look for job-related content in the body
          if len(body_text) > 200:
            job_description = body_text[:1500]  # Take first 1500 chars
            print(f"📝 Using page content as job description: {len(job_description)} characters")
        except:
          pass
      
      return job_description
      
    except Exception as e:
      print(f"⚠️ Error extracting job information: {e}")
      return ""
    
  def _detect_website_type(self):
    """Detect website type from URL to apply appropriate configuration"""
    if not self.link:
      return None
    
    url_lower = self.link.lower()
    for domain, config in self.website_configs.items():
      if domain in url_lower:
        print(f"🌐 Detected website type: {domain} ({config['type']})")
        return config
    
    print("🌐 Unknown website - using default configuration")
    return {
      "type": "unknown",
      "requires_account": False,
      "apply_button_required": False
    }
    
  def handle_apply_button(self):
    """Handle Apply button detection and clicking for websites that require it"""
    try:
      from selenium.webdriver.common.by import By
      from selenium.webdriver.support.ui import WebDriverWait
      from selenium.webdriver.support import expected_conditions as EC
      
      if not self.driver or not self.current_website_config:
        return True
      
      if not self.current_website_config.get("apply_button_required", False):
        return True
      
      print("🔍 Looking for Apply button...")
      
      # Common Apply button selectors
      apply_selectors = [
        "button:contains('Apply')",
        "a:contains('Apply')",
        "[data-qa*='apply']",
        "[href*='apply']",
        ".apply-button",
        "#apply-button",
        "input[value*='Apply']"
      ]
      
      # Website-specific selectors
      if "apply_button_selector" in self.current_website_config:
        apply_selectors.insert(0, self.current_website_config["apply_button_selector"])
      
      apply_button = None
      for selector in apply_selectors:
        try:
          if ":contains" in selector:
            # Use XPath for text-based selectors
            xpath_selector = f"//*[contains(text(), 'Apply') and (self::button or self::a)]"
            apply_button = self.driver.find_element(By.XPATH, xpath_selector)
          else:
            apply_button = self.driver.find_element(By.CSS_SELECTOR, selector)
          
          if apply_button and apply_button.is_displayed() and apply_button.is_enabled():
            break
        except:
          continue
      
      if apply_button:
        print("🎯 Found Apply button - clicking...")
        if not self.headless:
          self.driver.execute_script("arguments[0].style.border='3px solid blue'", apply_button)
          time.sleep(1)
        
        apply_button.click()
        print("✅ Apply button clicked")
        time.sleep(3)  # Wait for form to load
        return True
      else:
        print("⚠️ Apply button not found - continuing anyway")
        return True
        
    except Exception as e:
      print(f"⚠️ Error handling Apply button: {e}")
      return True  # Continue anyway
      
  def create_account_if_needed(self):
    """Create account if required by the website"""
    try:
      from selenium.webdriver.common.by import By
      
      if not self.driver or not self.current_website_config:
        return True
      
      if not self.current_website_config.get("requires_account", False):
        return True
      
      print("🔐 Account creation required for this website")
      
      # Look for account creation prompts
      account_indicators = [
        "sign up",
        "create account", 
        "register",
        "login required",
        "please sign in"
      ]
      
      page_text = ""
      try:
        page_text = self.driver.find_element(By.TAG_NAME, "body").text.lower()
      except:
        pass
      
      needs_account = any(indicator in page_text for indicator in account_indicators)
      
      if needs_account:
        print("🔑 Attempting to create account...")
        
        # Look for email field and sign up button
        email_selectors = [
          "input[type='email']",
          "input[name*='email']",
          "#email",
          ".email"
        ]
        
        signup_selectors = [
          "button:contains('Sign up')",
          "button:contains('Register')", 
          "button:contains('Create')",
          "[data-qa*='signup']",
          ".signup-button"
        ]
        
        email_field = None
        for selector in email_selectors:
          try:
            email_field = self.driver.find_element(By.CSS_SELECTOR, selector)
            if email_field.is_displayed():
              break
          except:
            continue
        
        if email_field:
          print("✉️ Found email field - entering email...")
          email_field.clear()
          email_field.send_keys(responses["email"])
          
          # Look for and click signup button
          for selector in signup_selectors:
            try:
              if ":contains" in selector:
                xpath_selector = f"//*[contains(text(), 'Sign up') or contains(text(), 'Register') or contains(text(), 'Create')]"
                signup_button = self.driver.find_element(By.XPATH, xpath_selector)
              else:
                signup_button = self.driver.find_element(By.CSS_SELECTOR, selector)
              
              if signup_button.is_displayed():
                signup_button.click()
                print("✅ Clicked signup button")
                time.sleep(3)
                break
            except:
              continue
        
        print("✅ Account creation process completed")
      
      return True
      
    except Exception as e:
      print(f"⚠️ Error with account creation: {e}")
      return True  # Continue anyway
    
  def customize_resume_with_placeholders(self, resume_path, company_name, job_title):
    """Customize resume by replacing <<Program>> and <<Company>> placeholders"""
    try:
      if not os.path.exists(resume_path):
        print(f"❌ Resume file not found: {resume_path}")
        return resume_path
      
      if not company_name and not job_title:
        print("⚠️ No company name or job title available for customization")
        return resume_path
      
      # Create customized filename
      base_name = os.path.splitext(resume_path)[0]
      extension = os.path.splitext(resume_path)[1]
      customized_path = f"{base_name}_{company_name.replace(' ', '')}{extension}"
      
      # For DOCX files, we need python-docx library
      try:
        from docx import Document
        
        # Load the document
        doc = Document(resume_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
          if "<<Program>>" in paragraph.text or "<<Company>>" in paragraph.text:
            print(f"📝 Found placeholders in paragraph: {paragraph.text[:100]}...")
            
            # Replace placeholders
            new_text = paragraph.text
            if job_title:
              new_text = new_text.replace("<<Program>>", job_title)
            if company_name:
              new_text = new_text.replace("<<Company>>", company_name)
            
            # Clear and set new text
            paragraph.clear()
            paragraph.add_run(new_text)
            print(f"✅ Updated to: {new_text[:100]}...")
        
        # Replace placeholders in tables (if any)
        for table in doc.tables:
          for row in table.rows:
            for cell in row.cells:
              if "<<Program>>" in cell.text or "<<Company>>" in cell.text:
                print(f"📝 Found placeholders in table cell: {cell.text[:50]}...")
                
                new_text = cell.text
                if job_title:
                  new_text = new_text.replace("<<Program>>", job_title)
                if company_name:
                  new_text = new_text.replace("<<Company>>", company_name)
                
                cell.text = new_text
                print(f"✅ Updated cell to: {new_text[:50]}...")
        
        # Save customized resume
        doc.save(customized_path)
        print(f"✅ Customized resume saved as: {customized_path}")
        return customized_path
        
      except ImportError:
        print("⚠️ python-docx not installed. Install with: pip install python-docx")
        print("⚠️ Using original resume without customization")
        return resume_path
      except Exception as e:
        print(f"❌ Error customizing resume: {e}")
        return resume_path
        
    except Exception as e:
      print(f"❌ Error in resume customization: {e}")
      return resume_path

  def analyze_job_and_select_resume(self, job_description=""):
    """Analyze job description and select appropriate resume"""
    try:
      from selenium.webdriver.common.by import By
      
      # Get job description from the page if not provided
      if not job_description and self.driver:
        try:
          # Look for job description on the page
          job_desc_selectors = [
            "[data-qa='job-description']",
            ".job-description",
            "#job-description", 
            ".description",
            ".content",
            "main",
            ".job-detail"
          ]
          
          for selector in job_desc_selectors:
            try:
              elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
              for element in elements:
                if element.text and len(element.text) > 200:
                  job_description = element.text[:2000]  # Limit text
                  break
              if job_description:
                break
            except:
              continue
                
        except Exception as e:
          print(f"⚠️ Could not extract job description: {e}")
      
      if not job_description:
        print("⚠️ No job description found, using default SWE resume")
        return self.resume_paths["swe"]
      
      # Analyze job type using LLM
      analysis_prompt = f"""Analyze the following job description and categorize it into one of these three types:

1. "swe" - Software Engineering roles (including quantitative software engineering, backend/frontend development, full-stack, DevOps, ML engineering, data engineering)
2. "quant" - Quantitative roles (research, analysis, strategy, trading, risk management, portfolio management, quantitative analysis)  
3. "communication" - Communication/business roles (receptionist, consulting, finance, marketing, sales, project management, business analyst)

Job Description:
{job_description[:1500]}

Respond with ONLY one word: either "swe", "quant", or "communication".
"""
      
      try:
        response = requests.post(
          f"{self.ollama_base_url}/api/generate",
          json={
            "model": self.model,
            "prompt": analysis_prompt,
            "stream": False,
            "options": {"temperature": 0.1}
          },
          timeout=30
        )
        
        if response.status_code == 200:
          result = response.json()
          job_type = result.get("response", "").strip().lower()
          
          if job_type in self.resume_paths:
            selected_resume = self.resume_paths[job_type]
            print(f"🤖 Job categorized as '{job_type}', selected resume: {selected_resume}")
            
            # Customize the selected resume with company and job details
            if self.company_name or self.job_title:
              print(f"🎯 Customizing resume for {self.company_name} - {self.job_title}")
              customized_resume = self.customize_resume_with_placeholders(
                selected_resume, 
                self.company_name, 
                self.job_title
              )
              return customized_resume
            
            return selected_resume
          
      except Exception as e:
        print(f"⚠️ LLM analysis failed: {e}")
      
      # Fallback to keyword-based analysis
      job_desc_lower = job_description.lower()
      
      # Check for quant keywords
      quant_keywords = ["quantitative", "quant", "research", "trading", "portfolio", "risk", "strategy", "analyst", "finance", "investment", "hedge fund", "alpha", "signals"]
      if any(keyword in job_desc_lower for keyword in quant_keywords):
        selected_resume = self.resume_paths["quant"]
        print(f"📊 Keyword analysis: Quant role detected, selected resume: {selected_resume}")
        
        # Customize the selected resume
        if self.company_name or self.job_title:
          print(f"🎯 Customizing resume for {self.company_name} - {self.job_title}")
          customized_resume = self.customize_resume_with_placeholders(
            selected_resume, 
            self.company_name, 
            self.job_title
          )
          return customized_resume
        
        return selected_resume
      
      # Check for communication keywords  
      comm_keywords = ["receptionist", "consulting", "consultant", "business", "marketing", "sales", "customer", "client", "communication", "project manager"]
      if any(keyword in job_desc_lower for keyword in comm_keywords):
        selected_resume = self.resume_paths["communication"]
        print(f"💬 Keyword analysis: Communication role detected, selected resume: {selected_resume}")
        
        # Customize the selected resume
        if self.company_name or self.job_title:
          print(f"🎯 Customizing resume for {self.company_name} - {self.job_title}")
          customized_resume = self.customize_resume_with_placeholders(
            selected_resume, 
            self.company_name, 
            self.job_title
          )
          return customized_resume
        
        return selected_resume
      
      # Default to SWE
      selected_resume = self.resume_paths["swe"]
      print(f"💻 Default: SWE role assumed, selected resume: {selected_resume}")
      
      # Customize the selected resume with company and job details
      if self.company_name or self.job_title:
        print(f"🎯 Customizing resume for {self.company_name} - {self.job_title}")
        customized_resume = self.customize_resume_with_placeholders(
          selected_resume, 
          self.company_name, 
          self.job_title
        )
        return customized_resume
      
      return selected_resume
      
    except Exception as e:
      print(f"❌ Error in job analysis: {e}")
      return self.resume_paths["swe"]

  def _load_profile_document(self):
    """Load profile document for enhanced LLM context"""
    profile_path = "Profile.txt"
    if os.path.exists(profile_path):
      try:
        with open(profile_path, 'r', encoding='utf-8') as f:
          content = f.read()
        print(f"📋 Loaded profile document: {profile_path}")
        return content
      except Exception as e:
        print(f"❌ Error loading profile document: {e}")
    else:
      print(f"⚠️ Profile document not found: {profile_path}")
    return ""

  def _load_reference_document(self):
    """Load reference document for RAG context"""
    if self.reference_doc_path and os.path.exists(self.reference_doc_path):
      try:
        with open(self.reference_doc_path, 'r', encoding='utf-8') as f:
          content = f.read()
        print(f"📄 Loaded reference document: {self.reference_doc_path}")
        return content
      except Exception as e:
        print(f"❌ Error loading reference document: {e}")
    return ""
    
  def _load_resume(self):
    """Load resume content for customization"""
    if self.resume_path and os.path.exists(self.resume_path):
      try:
        with open(self.resume_path, 'r', encoding='utf-8') as f:
          content = f.read()
        print(f"📄 Loaded resume: {self.resume_path}")
        return content
      except Exception as e:
        print(f"❌ Error loading resume: {e}")
    return ""
    
  def initialize_driver(self):
    """Initialize the web driver for browser automation"""
    try:
      from selenium import webdriver
      from selenium.webdriver.chrome.options import Options
      from webdriver_manager.chrome import ChromeDriverManager
      from selenium.webdriver.chrome.service import Service
      
      chrome_options = Options()
      chrome_options.add_argument("--no-sandbox")
      chrome_options.add_argument("--disable-dev-shm-usage")
      chrome_options.add_argument("--disable-gpu")
      chrome_options.add_argument("--remote-debugging-port=9222")
      
      # Only add headless mode if requested
      if self.headless:
        chrome_options.add_argument("--headless")
        print("🔍 Running in headless mode (no browser window)")
      else:
        print("👁️ Running with visible browser window (LIVE MODE)")
      
      service = Service(ChromeDriverManager().install())
      self.driver = webdriver.Chrome(service=service, options=chrome_options)
      return True
    except Exception as chrome_error:
      print(f"Chrome driver failed: {chrome_error}")
      
      # Try Firefox as fallback
      try:
        from selenium import webdriver
        from selenium.webdriver.firefox.options import Options as FirefoxOptions
        from webdriver_manager.firefox import GeckoDriverManager
        from selenium.webdriver.firefox.service import Service as FirefoxService
        
        firefox_options = FirefoxOptions()
        if self.headless:
          firefox_options.add_argument("--headless")
          print("🔍 Using Firefox in headless mode")
        else:
          print("👁️ Using Firefox with visible browser window (LIVE MODE)")
        
        service = FirefoxService(GeckoDriverManager().install())
        self.driver = webdriver.Firefox(service=service, options=firefox_options)
        print("🔄 Using Firefox driver as fallback")
        return True
      except Exception as firefox_error:
        print(f"Firefox driver also failed: {firefox_error}")
        print("Please install Chrome or Firefox browser")
        return False

  def get_dropdown_options(self, select_element):
    """Extract all available options from a dropdown element"""
    try:
      from selenium.webdriver.support.ui import Select
      select = Select(select_element)
      options = []
      for option in select.options:
        option_text = option.text.strip()
        option_value = option.get_attribute("value")
        if option_text and option_text.lower() not in ["select...", "choose...", "select one", "", "please select"]:
          options.append({
            "text": option_text,
            "value": option_value,
            "element": option
          })
      return options
    except Exception as e:
      print(f"❌ Error getting dropdown options: {e}")
      return []

  def find_best_dropdown_match(self, question_text, options, response_value):
    """Find the best matching option in a dropdown based on response value and question context"""
    if not options or not response_value:
      return None
    
    response_lower = response_value.lower()
    question_lower = question_text.lower()
    
    print(f"🔍 Matching '{response_value}' against {len(options)} options")
    
    # Exact matches first
    for option in options:
      if option["text"].lower() == response_lower:
        print(f"✅ Exact match found: {option['text']}")
        return option
    
    # Partial matches
    for option in options:
      option_text_lower = option["text"].lower()
      if response_lower in option_text_lower or option_text_lower in response_lower:
        print(f"✅ Partial match found: {option['text']}")
        return option
    
    # Context-based matching for specific question types
    if "sponsorship" in question_lower or "visa" in question_lower or "authorization" in question_lower:
      if response_value.lower() in ["yes", "true", "1"]:
        for option in options:
          if any(word in option["text"].lower() for word in ["yes", "required", "need", "true", "require"]):
            print(f"✅ Sponsorship 'Yes' match: {option['text']}")
            return option
      elif response_value.lower() in ["no", "false", "0"]:
        for option in options:
          if any(word in option["text"].lower() for word in ["no", "not required", "false", "don't", "authorized", "citizen"]):
            print(f"✅ Sponsorship 'No' match: {option['text']}")
            return option
    
    # For military service
    if "military" in question_lower:
      if response_value.lower() in ["no", "none"]:
        for option in options:
          if any(word in option["text"].lower() for word in ["no", "never", "not served", "none", "n/a"]):
            print(f"✅ Military service match: {option['text']}")
            return option
    
    # Generic Yes/No matching
    if response_value.lower() in ["yes", "no"]:
      for option in options:
        if option["text"].lower().startswith(response_value.lower()):
          print(f"✅ Yes/No match: {option['text']}")
          return option
    
    print(f"❌ No match found for '{response_value}' in dropdown options: {[opt['text'] for opt in options]}")
    return None

  def query_ollama(self, question_text, context="", custom_context=""):
    """Query Ollama LLM for generating responses to unknown questions"""
    try:
      # If it's a generic greeting or empty question, provide a default response
      if not question_text or question_text.lower().strip() in ["", "hello", "hi"]:
        return "I'm ready to assist you. What is the question?"
      
      # Extract company name from URL if not provided
      if not self.company_name and self.link:
        url_parts = self.link.lower()
        if "point72" in url_parts:
          self.company_name = "Point72"
        elif "goldman" in url_parts:
          self.company_name = "Goldman Sachs"
        elif "morgan" in url_parts:
          self.company_name = "J.P. Morgan"
        elif "blackrock" in url_parts:
          self.company_name = "BlackRock"
        elif "citadel" in url_parts:
          self.company_name = "Citadel"
      
      # Prepare the prompt with context and the actual question
      system_prompt = f"""You are helping fill out a job application form for Aditya Prabakaran (Adi), a qualified Computing student.

PROFILE INFORMATION:
{self.profile_content[:1500] if self.profile_content else 'Profile not available'}

KEY DETAILS:
- Name: Aditya Prabakaran (goes by "Adi")
- Education: Computer Science student at Imperial College London 
- Location: London, UK
- Email: aditya.prabakaran@gmail.com
- LinkedIn: https://www.linkedin.com/in/adiprabs
- GitHub: https://github.com/TheDarkEyezor
- Portfolio: https://adiprabs.vercel.app/
- US Work Authorization: Requires sponsorship
- UK Work Authorization: No sponsorship needed
- Military Service: None

{f'Company: {self.company_name}' if self.company_name else ''}
{f'Job Title: {self.job_title}' if self.job_title else ''}

{f'Resume Context: {self.resume_content[:1000]}...' if self.resume_content else ''}
{f'Reference Document: {context[:500]}...' if context else ''}
{f'Additional Context: {custom_context}' if custom_context else ''}

Answer the following job application question concisely and professionally. 
Use the profile information above to provide accurate, personalized responses.
If it's a yes/no question, respond with just "Yes" or "No".
If it's asking for specific information, provide a brief, relevant answer.
For cover letter or essay questions, provide 2-3 sentences highlighting relevant experience from the profile.
If you need to make assumptions, make reasonable ones based on the profile above.

IMPORTANT: This is the actual question from the form: "{question_text}"

Answer:"""

      payload = {
        "model": self.model,
        "prompt": system_prompt,
        "stream": False,
        "options": {
          "temperature": 0.1,  # Low temperature for consistent answers
          "top_p": 0.9
        }
      }
      
      response = requests.post(
        f"{self.ollama_base_url}/api/generate",
        json=payload,
        timeout=30
      )
      
      if response.status_code == 200:
        result = response.json()
        answer = result.get("response", "").strip()
        
        # Check if we got a valid response (not a generic "ready to assist" message)
        if answer and not answer.lower().startswith("i'm ready to"):
          print(f"🤖 Ollama generated response for '{question_text[:50]}...': {answer[:100]}...")
          return answer
        else:
          print(f"⚠️ LLM gave generic response, using fallback")
          return self.get_fallback_response(question_text)
      else:
        print(f"❌ Ollama API error: {response.status_code}")
        return self.get_fallback_response(question_text)
        
    except requests.exceptions.ConnectionError:
      print("❌ Could not connect to Ollama. Is it running? Start with: ollama serve")
      return self.get_fallback_response(question_text)
    except Exception as e:
      print(f"❌ Error querying Ollama: {e}")
      return self.get_fallback_response(question_text)

  def get_fallback_response(self, question):
    """Generate a fallback response when LLM fails"""
    question_lower = question.lower()
    
    if any(word in question_lower for word in ["name", "full name"]):
      return "Aditya Prabakaran"
    elif any(word in question_lower for word in ["email", "e-mail"]):
      return "aditya.prabakaran@gmail.com"
    elif any(word in question_lower for word in ["phone", "mobile", "telephone"]):
      return "+44 7700 123456"
    elif any(word in question_lower for word in ["address", "location"]):
      return "London, UK"
    elif any(word in question_lower for word in ["linkedin", "portfolio", "website"]):
      return "https://linkedin.com/in/adiprabs"
    elif any(word in question_lower for word in ["cover letter", "why", "interest", "motivation"]):
      return "I am very interested in this position due to my strong background in computer science and software engineering."
    else:
      return "Yes"

  def generate_customized_resume(self, job_description=""):
    """Generate a customized resume/cover letter for the specific job"""
    if not self.resume_content:
      return ""
    
    context = f"""
    Job Description: {job_description}
    Company: {self.company_name}
    Position: {self.job_title}
    """
    
    prompt = f"""Based on the following job requirements, customize this resume/cover letter to highlight the most relevant experience and skills.

Original Resume:
{self.resume_content}

{context}

Provide a customized version that emphasizes relevant experience for this specific role. Keep it professional and concise."""
    
    return self.query_ollama("Customize resume for this position", custom_context=prompt)

  def handle_file_upload(self, element, field_text):
    """Handle file upload fields, specifically for resume/CV. For cover letters, prefer manual entry."""
    try:
      field_lower = field_text.lower()
      
      # Check if this is a cover letter field - prefer manual entry
      if any(keyword in field_lower for keyword in ["cover letter", "covering letter", "motivation letter", "personal statement"]):
        print(f"📝 Detected cover letter field: {field_text}")
        print("💬 Cover letters should be entered manually in text areas, not as file uploads")
        
        # Look for a manual entry option or text area nearby
        try:
          from selenium.webdriver.common.by import By
          if self.driver:
            # Try to find "Enter manually" or similar options
            manual_selectors = [
              "//button[contains(text(), 'Enter manually')]",
              "//button[contains(text(), 'Type manually')]", 
              "//a[contains(text(), 'Enter manually')]",
              "//span[contains(text(), 'manual')]",
              "//div[contains(@class, 'manual') and contains(text(), 'Enter')]"
            ]
            
            for selector in manual_selectors:
              try:
                manual_button = self.driver.find_element(By.XPATH, selector)
                if manual_button.is_displayed():
                  manual_button.click()
                  print(f"✅ Clicked 'Enter manually' option for cover letter")
                  time.sleep(2)
                  return True
              except:
                continue
                
        except Exception as e:
          print(f"⚠️ Could not find manual entry option: {e}")
        
        print("⚠️ Cover letter field detected but no manual entry option found")
        return False
      
      # Check if this is a resume/CV upload field
      elif any(keyword in field_lower for keyword in ["resume", "cv", "curriculum", "upload", "attach", "file"]):
        if self.resume_path and os.path.exists(self.resume_path):
          print(f"📎 Attempting to upload resume: {self.resume_path}")
          
          # Small wait before uploading to avoid interrupting active widgets
          time.sleep(0.4)

          # Highlight file upload field
          if not self.headless and self.driver:
            self.driver.execute_script("arguments[0].style.border='3px solid purple'", element)
            time.sleep(1)
          
          # For Greenhouse forms, we might need to handle file upload differently
          try:
            # Method 1: Direct file input
            element.send_keys(os.path.abspath(self.resume_path))
            print(f"✅ Resume uploaded via direct input")

            # Small stability wait after upload to let the page register the file
            time.sleep(0.6)

            # Try to verify uploaded filename appears nearby (common in Greenhouse)
            try:
              nearby_text = element.find_element_by_xpath("..")
              if nearby_text and (os.path.basename(self.resume_path) in nearby_text.text):
                print("✅ Upload verified by nearby filename text")
            except:
              pass
            
          except Exception as e1:
            print(f"⚠️ Direct upload failed, trying alternative method: {e1}")
            try:
              # Method 2: Click and then send keys
              element.click()
              time.sleep(1)
              element.send_keys(os.path.abspath(self.resume_path))
              print(f"✅ Resume uploaded via click-then-send")
              
            except Exception as e2:
              print(f"⚠️ Click upload failed, trying manual entry: {e2}")
              try:
                # Try to find the "Enter manually" option instead
                from selenium.webdriver.common.by import By
                if self.driver:
                  manual_selectors = [
                    "//button[contains(text(), 'Enter manually')]",
                    "//button[contains(text(), 'manual')]", 
                    "//a[contains(text(), 'Enter manually')]",
                    "//span[contains(text(), 'Enter manually')]",
                    "//*[contains(@class, 'manual') and contains(text(), 'Enter')]"
                  ]
                  
                  for selector in manual_selectors:
                    try:
                      manual_button = self.driver.find_element(By.XPATH, selector)
                      if manual_button.is_displayed():
                        manual_button.click()
                        print(f"✅ Clicked 'Enter manually' button as fallback")
                        time.sleep(2)
                        return True
                    except:
                      continue
                      
              except Exception as e3:
                print(f"❌ Could not find manual entry option: {e3}")
                return False
          
          # Try to refresh and find file input again if stale element
          try:
            if not self.headless and self.driver:
              time.sleep(1)
              self.driver.execute_script("arguments[0].style.border='3px solid green'", element)
              print(f"✅ Resume upload process completed")
          except:
            print(f"⚠️ Could not provide visual feedback - element may be stale")
          
          return True
        else:
          print(f"❌ Resume path not provided or file not found: {self.resume_path}")
          return False
      
      return False
    except Exception as e:
      print(f"❌ Error handling file upload: {e}")
      return False
      
  def find_matching_response(self, field_text):
    """Find matching response from the responses dictionary"""
    field_text_lower = field_text.lower()
    
    # First check for website-specific responses
    website_response = self.get_website_specific_response(field_text)
    if website_response:
      return website_response
    
    # Direct mapping for common fields
    field_mappings = {
      "first name": "first name",
      "last name": "last name", 
      "full name": "full name",
      "email": "email",
      "phone": "phone",
      "location": "city",
      "city": "city",
      "country": "country",
      "school": "school",
      "degree": "degree",
      "major": "major",
      "linkedin": "linkedin",
      "github": "github",
      "profile": "profile"
    }
    
    # Check for direct matches
    for key, response_key in field_mappings.items():
      if key in field_text_lower:
        return responses.get(response_key, "")
    
    # Check for sponsorship questions
    if "sponsorship" in field_text_lower and "united states" in field_text_lower:
      return "Yes" if responses.get("united states need sponsorship") == "yes" else "No"
    
    if "sponsorship" in field_text_lower and "united kingdom" in field_text_lower:
      return "Yes" if responses.get("united kingdom need sponsorship") == "yes" else "No"
    
    # General sponsorship questions
    if "sponsorship" in field_text_lower or "visa" in field_text_lower:
      return "Yes" if responses.get("united states need sponsorship") == "yes" else "No"
    
    # Check for military service
    if "military" in field_text_lower:
      return "No" if responses.get("military service") == "none" else "Yes"
    
    # Check for work authorization
    if "authorized" in field_text_lower and "work" in field_text_lower:
      if "united states" in field_text_lower:
        return "No" if responses.get("united states need sponsorship") == "yes" else "Yes"
      return "Yes"  # Default for other countries
    
    # Check for privacy/consent questions
    if "privacy" in field_text_lower or "consent" in field_text_lower:
      return "Yes"
    
    # Check for "I accept" or similar checkbox/agreement fields
    if any(phrase in field_text_lower for phrase in ["i accept", "accept", "agree", "consent"]):
      return "Yes"
    
    # Check for previous application
    if "previously applied" in field_text_lower:
      return "No"  # Default assumption
    
    return None
  
  def get_website_specific_response(self, field_text):
    """Get website-specific responses for known fields"""
    if not self.current_website_config or "specific_fields" not in self.current_website_config:
      return None
    
    field_lower = field_text.lower()
    specific_fields = self.current_website_config["specific_fields"]
    
    # Check for specific field mappings
    for key, value in specific_fields.items():
      if key.replace("_", " ") in field_lower:
        print(f"🎯 Website-specific response for '{field_text}': {value}")
        return value
    
    # Special handling for common patterns
    if "pronouns" in field_lower:
      return specific_fields.get("pronouns", "He/him")
    elif "formula" in field_lower and "experience" in field_lower:
      return specific_fields.get("formula1_experience", "No")
    elif "previous" in field_lower and "employee" in field_lower:
      return specific_fields.get("previous_revolut_employee", "No")
    elif "interview" in field_lower and "transcript" in field_lower:
      return specific_fields.get("interview_transcript_consent", "Yes, I consent")
    
    return None
  
  def get_field_label(self, element):
    """Extract the label or question text for a form field"""
    try:
      from selenium.webdriver.common.by import By
      
      # Try multiple methods to find the field label
      label_text = ""
      
      try:
        # Try to find associated label
        label_element = element.find_element(By.XPATH, ".//preceding::label[1]")
        label_text = label_element.text
      except:
        try:
          # Try aria-labelledby
          aria_labelledby = element.get_attribute("aria-labelledby")
          if aria_labelledby and self.driver:
            label_element = self.driver.find_element(By.ID, aria_labelledby)
            label_text = label_element.text
        except:
          try:
            # Try placeholder
            label_text = element.get_attribute("placeholder") or ""
          except:
            try:
              # Try aria-label
              label_text = element.get_attribute("aria-label") or ""
            except:
              try:
                # Try nearby text or name
                label_text = element.get_attribute("name") or ""
              except:
                try:
                  # Try parent element text
                  parent = element.find_element(By.XPATH, "./..")
                  label_text = parent.text[:50] if parent.text else ""
                except:
                  pass
      
      return label_text.strip() if label_text else ""
      
    except Exception as e:
      print(f"❌ Error getting field label: {e}")
      return ""

  def highlight_element(self, element):
    """Highlight an element for visual feedback"""
    try:
      if not self.headless and self.driver:
        original_style = element.get_attribute("style") or ""
        self.driver.execute_script("arguments[0].style.border='3px solid red'", element)
    except Exception as e:
      print(f"⚠️ Could not highlight element: {e}")

  def fill_custom_dropdown(self, dropdown_element):
    """Fill Greenhouse-style custom dropdown element (div with role='combobox').

    Strategy:
    - Try to type into the internal input (react-select style), press Enter, then Tab.
    - If typing fails, fall back to opening the menu, clicking the matching option,
      pressing Enter, and Tabbing to the next field.
    - Verify the selection by checking the displayed single-value element.
    """
    try:
      from selenium.webdriver.common.by import By
      from selenium.webdriver.support.ui import WebDriverWait
      from selenium.webdriver.support import expected_conditions as EC
      from selenium.webdriver.common.keys import Keys

      if not self.driver:
        return False

      # Get label/question context for this dropdown
      question = self.get_field_label(dropdown_element)
      if not question or len(question.strip()) < 3:
        try:
          label_element = dropdown_element.find_element(By.TAG_NAME, "label")
          question = label_element.text
        except:
          return False

      # Use predefined responses first, then LLM
      response_value = self.find_matching_response(question)
      if response_value is None:
        response_value = self.query_ollama(question)

      if not response_value or "ready to assist" in response_value.lower():
        print(f"⚠️ Skipping custom dropdown - no valid response for: {question}")
        return False

      if not self.headless:
        self.highlight_element(dropdown_element)

      print(f"🔽 Attempting to fill custom dropdown: '{question}' with '{response_value}'")

      # 1. Find the dropdown control and open it
      control_element = dropdown_element.find_element(By.CSS_SELECTOR, "div[class*='select__control']")
      self.driver.execute_script("arguments[0].click();", control_element)

      wait = WebDriverWait(self.driver, 5)
      menu_selector = "div[class*='select__menu']"

      # 2. Preferred method: type into internal input, press Enter, then Tab
      try:
        input_selector = "input.select__input, input[class*='select__input']"
        input_elem = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, input_selector)))

        try:
          input_elem.clear()
        except:
          pass
        input_elem.send_keys(response_value)
        time.sleep(0.05)
        input_elem.send_keys(Keys.ENTER)
        # Dispatch additional events to ensure React/JS listeners notice the change
        try:
          self.driver.execute_script(
            "var el = arguments[0]; el.dispatchEvent(new Event('input', {bubbles: true})); el.dispatchEvent(new Event('change', {bubbles: true})); el.dispatchEvent(new KeyboardEvent('keydown', {key:'Enter', bubbles:true})); el.dispatchEvent(new KeyboardEvent('keyup', {key:'Enter', bubbles:true})); el.blur();",
            input_elem
          )
        except Exception as e_events:
          print(f"⚠️ Could not dispatch JS events after typing: {e_events}")
        print("✅ Typed response into input and pressed Enter to confirm.")

        try:
          input_elem.send_keys(Keys.TAB)
          print("✅ Sent TAB to move focus to next field.")
        except Exception:
          try:
            self.driver.execute_script("var e = arguments[0]; var next = e.parentElement.querySelector('input, textarea'); if(next) next.focus();", input_elem)
          except:
            pass

        try:
          wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, menu_selector)))
        except:
          pass

        try:
          selected_text_element = dropdown_element.find_element(By.CSS_SELECTOR, "div[class*='select__single-value']")
          selected_text = selected_text_element.text
          if response_value.strip().lower() in selected_text.lower():
            print(f"✅ Verification successful after Enter+Tab. Dropdown shows: '{selected_text}'")
            if not self.headless:
              self.driver.execute_script("arguments[0].style.border='3px solid green'", dropdown_element)
              time.sleep(0.3)
            return True
          else:
            print(f"⚠️ Verification mismatch. Expected '{response_value}', got '{selected_text}'")
        except Exception:
          print("⚠️ Could not verify selected value after Enter+Tab; will try fallback click method.")

      except Exception as e_input:
        print(f"⚠️ Typing method failed or input not found: {e_input}")

      # 3. Fallback: open menu, click matching option, press Enter, Tab
      try:
        menu = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, menu_selector)))
        print("✅ Dropdown menu is visible (fallback path).")

        response_text = response_value.strip()
        option_xpath = f"//div[contains(@class, 'select__option') and normalize-space(translate(., 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'))='{response_text.lower()}']"
        target_option = wait.until(EC.visibility_of_element_located((By.XPATH, option_xpath)))
        print(f"✅ Found option (fallback): '{target_option.text}'")

        self.driver.execute_script("arguments[0].click();", target_option)
        try:
          control_element.send_keys(Keys.ENTER)
          print("✅ Sent 'Enter' key to confirm selection (fallback).")
        except Exception as e:
          print(f"⚠️ Could not send Enter on fallback: {e}")

        try:
          control_element.send_keys(Keys.TAB)
          time.sleep(0.1)
        except:
          pass

        try:
          wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, menu_selector)))
        except:
          pass

        try:
          selected_text_element = dropdown_element.find_element(By.CSS_SELECTOR, "div[class*='select__single-value']")
          selected_text = selected_text_element.text
          if response_text.lower() in selected_text.lower():
            print(f"✅ Verification successful (fallback). Dropdown shows: '{selected_text}'")
            if not self.headless:
              self.driver.execute_script("arguments[0].style.border='3px solid green'", dropdown_element)
              time.sleep(0.3)
            return True
          else:
            print(f"⚠️ Verification failed (fallback). Expected '{response_text}', but dropdown shows '{selected_text}'")
            return False
        except Exception:
          print(f"⚠️ Could not verify selection after fallback click; assuming success.")
          return True

      except Exception as e_fallback:
        print(f"❌ Fallback click method failed too: {e_fallback}")
        try:
          self.driver.find_element(By.TAG_NAME, 'body').click()
        except:
          pass
        return False

    except Exception as e:
      print(f"❌ Error filling custom dropdown: {e}")
      return False

  def fill_response(self, element):
    """Fill a form element with appropriate response using enhanced logic"""
    try:
      from selenium.webdriver.common.by import By
      from selenium.webdriver.support.ui import Select
      from selenium.webdriver.support.ui import WebDriverWait
      from selenium.webdriver.support import expected_conditions as EC
      from selenium.webdriver.common.keys import Keys
      
      if not self.driver:
        print("❌ Driver not initialized")
        return False
      
      # Get the label or placeholder text to determine what to fill
      label_text = ""
      
      # Try multiple methods to find the field label
      try:
        # Try to find associated label
        label_element = element.find_element(By.XPATH, ".//preceding::label[1]")
        label_text = label_element.text
      except:
        try:
          # Try aria-labelledby
          aria_labelledby = element.get_attribute("aria-labelledby")
          if aria_labelledby and self.driver:
            label_element = self.driver.find_element(By.ID, aria_labelledby)
            label_text = label_element.text
        except:
          try:
            # Try placeholder
            label_text = element.get_attribute("placeholder") or ""
          except:
            try:
              # Try aria-label
              label_text = element.get_attribute("aria-label") or ""
            except:
              try:
                # Try nearby text or name
                label_text = element.get_attribute("name") or ""
              except:
                try:
                  # Try parent element text
                  parent = element.find_element(By.XPATH, "./..")
                  label_text = parent.text[:50] if parent.text else ""
                except:
                  pass
      
      if not label_text:
        print("⚠️ Could not determine field label")
        return False
      
      # Handle file upload fields first
      input_type = element.get_attribute("type")
      if input_type == "file":
        return self.handle_file_upload(element, label_text)
      
      # Find matching response from predefined responses
      response_value = self.find_matching_response(label_text)
      
      # If no predefined response found, use Ollama LLM
      if response_value is None:
        print(f"🤖 Using LLM for unknown question: {label_text}")
        response_value = self.query_ollama(label_text, self.reference_content)
      
      if response_value:
        # Highlight the field before filling (visual feedback)
        original_style = ""
        if not self.headless and self.driver:
          original_style = element.get_attribute("style") or ""
          self.driver.execute_script("arguments[0].style.border='3px solid red'", element)
          print(f"🔍 Filling field '{label_text}' with: {response_value}")
        
        # Add delay for visual effect if in slow mode
        if self.slow_mode and not self.headless:
          time.sleep(1)
        
        # Handle different input types
        tag_name = element.tag_name.lower()
        
        if tag_name == "select":
          # Enhanced dropdown handling with proper selection
          options = self.get_dropdown_options(element)
          if options:
            print(f"📋 Found {len(options)} dropdown options")
            
            best_match = self.find_best_dropdown_match(label_text, options, response_value)
            if best_match:
              try:
                from selenium.webdriver.common.keys import Keys
                
                # Click the dropdown to open it
                element.click()
                time.sleep(0.5)
                
                select = Select(element)
                
                # Try different selection methods
                selection_success = False
                try:
                  select.select_by_visible_text(best_match["text"])
                  print(f"✅ Selected dropdown option: {best_match['text']}")
                  selection_success = True
                except:
                  try:
                    select.select_by_value(best_match["value"])
                    print(f"✅ Selected dropdown option by value: {best_match['value']}")
                    selection_success = True
                  except:
                    try:
                      # Try clicking the option directly
                      best_match["element"].click()
                      print(f"✅ Clicked dropdown option: {best_match['text']}")
                      selection_success = True
                    except:
                      print(f"❌ Could not select dropdown option: {best_match['text']}")
                
                # Press Enter to confirm selection for Greenhouse dropdowns
                if selection_success:
                  time.sleep(0.5)
                  # Try pressing Enter on both the select element and the selected option
                  try:
                    element.send_keys(Keys.ENTER)
                    print(f"✅ Pressed Enter on select element to confirm selection")
                  except:
                    try:
                      best_match["element"].send_keys(Keys.ENTER)
                      print(f"✅ Pressed Enter on selected option to confirm")
                    except:
                      print(f"⚠️ Could not press Enter to confirm selection")
                  time.sleep(0.5)
                      
                # Verify selection worked
                try:
                  selected_option = select.first_selected_option.text
                  if selected_option != best_match["text"]:
                    print(f"⚠️ Selection may not have worked. Expected: {best_match['text']}, Got: {selected_option}")
                  else:
                    print(f"✅ Dropdown selection verified: {selected_option}")
                except:
                  print("⚠️ Could not verify dropdown selection")
                
              except Exception as e:
                print(f"❌ Error selecting dropdown: {e}")
            else:
              print(f"❌ No matching dropdown option found for: {response_value}")
              print(f"Available options: {[opt['text'] for opt in options]}")
          
        elif tag_name == "input" and input_type in ["text", "email", "tel"]:
          element.clear()

          # Determine if this input behaves like a combobox (react-select or aria autocompletion)
          role_attr = (element.get_attribute("role") or "").lower()
          aria_auto = (element.get_attribute("aria-autocomplete") or "").lower()
          class_attr = (element.get_attribute("class") or "").lower()
          is_combobox_like = ("combobox" in role_attr) or (aria_auto == "list") or ("select__input" in class_attr)

          if self.slow_mode and not self.headless:
            # Type character by character for visual effect
            for char in response_value:
              element.send_keys(char)
              time.sleep(0.1)
          else:
            element.send_keys(response_value)

          # If this behaves like a combobox, press Enter immediately to confirm and Tab to move on
          if is_combobox_like:
            try:
              element.send_keys(Keys.ENTER)
              time.sleep(0.05)
              element.send_keys(Keys.TAB)
              # Click outside to force blur and propagation
              try:
                self.driver.find_element(By.TAG_NAME, 'body').click()
              except:
                pass
              time.sleep(0.12)
              print(f"✅ Sent ENTER+TAB for combobox-like input to confirm selection")
              # Try to verify the selection appeared nearby; if not, attempt robust dropdown fill
              try:
                # Look for a nearby single-value display used by react-select
                sel_text = None
                try:
                  container = element.find_element(By.XPATH, "ancestor::div[contains(@class, 'select')][1]")
                except:
                  container = None
                if container:
                  try:
                    sel_el = container.find_element(By.CSS_SELECTOR, "div[class*='select__single-value']")
                    sel_text = sel_el.text
                  except:
                    try:
                      ctrl = container.find_element(By.CSS_SELECTOR, "div[class*='select__control']")
                      sel_text = ctrl.text
                    except:
                      sel_text = None

                if sel_text and response_value.strip().lower() not in sel_text.lower():
                  print(f"⚠️ Combobox verification mismatch: expected '{response_value}', got '{sel_text}', retrying via custom dropdown handler")
                  try:
                    combo = element.find_element(By.XPATH, "ancestor::div[@role='combobox'][1]")
                  except:
                    combo = container
                  if combo:
                    self.fill_custom_dropdown(combo)
              except Exception:
                pass
            except Exception as e:
              print(f"⚠️ Could not send ENTER+TAB for combobox input: {e}")
            
        elif tag_name == "textarea":
          element.clear()
          
          # Get the actual question text for LLM processing
          actual_question = label_text
          
          # Try to get more context from the form
          try:
            # Look for associated labels or nearby text that might contain the question
            parent = element.find_element(By.XPATH, "./..")
            parent_text = parent.text.strip()
            if parent_text and len(parent_text) > len(label_text):
              actual_question = parent_text
              print(f"📝 Found extended question context: {actual_question[:100]}...")
          except:
            pass
          
          # For textarea, check if it needs customized content
          if any(keyword in actual_question.lower() for keyword in ["cover letter", "why", "interest", "motivation", "essay", "note", "hiring manager"]):
            print(f"🤖 Generating customized response for textarea question: {actual_question[:100]}...")
            # Use LLM to generate customized response with the actual question
            customized_response = self.query_ollama(actual_question, self.reference_content, f"This is for {self.company_name}")
            if customized_response and not customized_response.startswith("I'm ready to assist"):
              response_value = customized_response
              print(f"✅ Generated customized response: {response_value[:100]}...")
            else:
              # Fallback to a professional default response
              response_value = f"I am very interested in this {self.job_title} position at {self.company_name}. With my MEng in Computer Science from Imperial College London and experience in software engineering, I believe I would be a strong addition to your team."
          
          if self.slow_mode and not self.headless:
            # Type character by character for visual effect
            for char in response_value:
              element.send_keys(char)
              time.sleep(0.05)
          else:
            element.send_keys(response_value)
        
        # Restore original border style and add completion indicator
        if not self.headless and self.driver:
          time.sleep(0.5)  # Brief pause to show completion
          self.driver.execute_script("arguments[0].style.border='3px solid green'", element)
          time.sleep(0.5)
          self.driver.execute_script(f"arguments[0].style='{original_style}'", element)
          print(f"✅ Completed field: {label_text}")
        
        return True
      else:
        print(f"⚠️ No response generated for field: {label_text}")
      
    except Exception as e:
      print(f"❌ Error filling element: {e}")
      return False
    
    return False
  
  def fill_form(self):
    """Fill the entire form on the current page"""
    try:
      from selenium.webdriver.common.by import By
      
      if not self.driver:
        if not self.initialize_driver():
          return False
      
      if not self.driver:
        print("❌ Failed to initialize driver")
        return False
      
      print(f"🌐 Navigating to: {self.link}")
      self.driver.get(self.link)
      
      # Add a small delay to let the page load
      time.sleep(3)
      
      print("🔍 Searching for form fields...")
      
      # Find all form elements including custom dropdowns
      input_fields = self.driver.find_elements(By.TAG_NAME, "input")
      select_fields = self.driver.find_elements(By.TAG_NAME, "select")
      textarea_fields = self.driver.find_elements(By.TAG_NAME, "textarea")
      
      # Look for custom dropdowns (common in modern web forms)
      custom_dropdowns = []
      try:
        # More specific selectors for actual Greenhouse dropdowns
        dropdown_selectors = [
          "div[role='combobox']",
          "div[class*='select'][role='button']",
          "button[class*='select']",
          "div[class*='dropdown'][role='button']",
          ".input_wrapper div[role='button']",
          ".field div[role='button']"
        ]
        
        for selector in dropdown_selectors:
          elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
          for element in elements:
            if element.is_displayed() and element not in custom_dropdowns:
              # Additional filter: must have some meaningful text or aria-label
              label = self.get_field_label(element)
              if label and len(label.strip()) > 3:
                custom_dropdowns.append(element)
              
        # Limit to reasonable number to avoid processing too many elements
        if len(custom_dropdowns) > 10:
          print(f"⚠️ Found {len(custom_dropdowns)} custom dropdowns, limiting to first 10")
          custom_dropdowns = custom_dropdowns[:10]
          
      except Exception as e:
        print(f"⚠️ Error finding custom dropdowns: {e}")
      
      filled_count = 0
      # Collect file inputs separately if we should defer uploads
      defer_uploads = False
      try:
        defer_uploads = bool(self.current_website_config.get("defer_file_uploads", False)) if self.current_website_config else False
      except:
        defer_uploads = False

      deferred_file_inputs = []

      print(f"📊 Found {len(input_fields)} input fields, {len(select_fields)} select fields, {len(textarea_fields)} textarea fields, {len(custom_dropdowns)} custom dropdowns")
      
      # Fill input fields (including file uploads)
      for i, field in enumerate(input_fields):
        try:
          # Re-find the element to avoid stale reference
          try:
            field = self.driver.find_elements(By.TAG_NAME, "input")[i]
          except:
            continue
            
          field_type = field.get_attribute("type")
          # If deferring uploads, collect file inputs to process later
          if field_type == "file" and defer_uploads:
            if field.is_displayed():
              print(f"📂 Deferring file input {i+1}/{len(input_fields)} to later")
              deferred_file_inputs.append((i, field))
            continue

          if field.is_displayed() and field_type in ["text", "email", "tel", "file"]:
            print(f"\n🔄 Processing input field {i+1}/{len(input_fields)} (type: {field_type})")
            if self.fill_response(field):
              filled_count += 1
            if self.slow_mode and not self.headless:
              time.sleep(0.8)  # Pause between fields
        except Exception as e:
          print(f"❌ Error with input field {i+1}: {e}")
      
      # Fill select fields (traditional dropdowns)
      for i, field in enumerate(select_fields):
        try:
          # Re-find the element to avoid stale reference
          try:
            field = self.driver.find_elements(By.TAG_NAME, "select")[i]
          except:
            continue
            
          if field.is_displayed():
            print(f"\n🔄 Processing dropdown field {i+1}/{len(select_fields)}")
            if self.fill_response(field):
              filled_count += 1
            if self.slow_mode and not self.headless:
              time.sleep(0.8)
        except Exception as e:
          print(f"❌ Error with select field {i+1}: {e}")
      
      # Fill custom dropdowns
      for i, field in enumerate(custom_dropdowns):
        try:
          if field.is_displayed():
            print(f"\n🔄 Processing custom dropdown {i+1}/{len(custom_dropdowns)}")
            if self.fill_custom_dropdown(field):
              filled_count += 1
            if self.slow_mode and not self.headless:
              time.sleep(0.8)
        except Exception as e:
          print(f"❌ Error with custom dropdown {i+1}: {e}")
      
      # Fill textarea fields
      for i, field in enumerate(textarea_fields):
        try:
          # Re-find the element to avoid stale reference
          try:
            field = self.driver.find_elements(By.TAG_NAME, "textarea")[i]
          except:
            continue
            
          if field.is_displayed():
            print(f"\n🔄 Processing textarea field {i+1}/{len(textarea_fields)}")
            if self.fill_response(field):
              filled_count += 1
            if self.slow_mode and not self.headless:
              time.sleep(0.8)
        except Exception as e:
          print(f"❌ Error with textarea field {i+1}: {e}")
      
      print(f"\n✅ Successfully filled {filled_count} fields")

      # Now process deferred file uploads (if any)
      if deferred_file_inputs:
        print(f"\n📂 Processing {len(deferred_file_inputs)} deferred file upload(s) now...")
        for idx, file_field in deferred_file_inputs:
          try:
            print(f"🔄 Uploading deferred file for field index {idx}")
            if self.handle_file_upload(file_field, self.get_field_label(file_field)):
              filled_count += 1
              print(f"✅ Deferred file upload succeeded for field index {idx}")
            else:
              print(f"⚠️ Deferred file upload failed for field index {idx}")
            time.sleep(0.5)
          except Exception as e:
            print(f"⚠️ Error during deferred file upload for field {idx}: {e}")

      
      if not self.headless:
        print("⏸️  Pausing for 5 seconds to review filled form...")
        time.sleep(5)
      return True
    except Exception as e:
      print(f"❌ Error filling form: {e}")
      return False

      return True
      
    except Exception as e:
      print(f"❌ Error filling form: {e}")
      return False
  
  def report_and_refill_missing_fields(self):
    """Find fields that are still invalid/required and attempt to refill them."""
    try:
      from selenium.webdriver.common.by import By
      missing = []

      # 1) Find elements marked aria-invalid
      try:
        invalids = self.driver.find_elements(By.CSS_SELECTOR, "[aria-invalid='true']")
        for el in invalids:
          try:
            label = self.get_field_label(el) or el.get_attribute('name') or el.get_attribute('id') or el.text
            missing.append((label, el))
          except:
            continue
      except:
        pass

      # 2) Find visible error messages saying 'This field is required'
      try:
        bodies = self.driver.find_elements(By.XPATH, "//*[contains(text(), 'This field is required') or contains(text(), 'required')]")
        for b in bodies:
          try:
            # Attempt to find the nearest input/select/textarea
            parent = b.find_element(By.XPATH, "./ancestor::label[1] | ./ancestor::div[1]")
            try:
              candidate = parent.find_element(By.CSS_SELECTOR, "input, select, textarea, div[role='combobox']")
              label = self.get_field_label(candidate) or candidate.get_attribute('name') or candidate.get_attribute('id') or candidate.text
              missing.append((label, candidate))
            except:
              continue
          except:
            continue
      except:
        pass

      # Deduplicate
      seen = set()
      unique_missing = []
      for label, el in missing:
        key = (label or '') + '|' + (el.get_attribute('outerHTML')[:200] if el else '')
        if key not in seen:
          seen.add(key)
          unique_missing.append((label, el))

      if not unique_missing:
        print("🔎 No missing/invalid fields detected after fill.")
        return True

      print(f"🔄 Found {len(unique_missing)} missing/invalid fields. Attempting to refill...")
      for label, el in unique_missing:
        try:
          print(f"↺ Refilling: {label}")
          tag = el.tag_name.lower() if el else ''
          if tag == 'input' or tag == 'textarea' or tag == 'select':
            self.fill_response(el)
          else:
            # Try custom dropdown handler
            try:
              self.fill_custom_dropdown(el)
            except:
              # as last resort, try clicking then tabbing
              try:
                el.click()
                from selenium.webdriver.common.keys import Keys
                el.send_keys(Keys.ENTER)
                el.send_keys(Keys.TAB)
              except:
                pass
          time.sleep(0.3)
        except Exception as e:
          print(f"⚠️ Could not refill {label}: {e}")

      return True
    except Exception as e:
      print(f"❌ Error in report_and_refill_missing_fields: {e}")
      return False

  def run_multiple_applications(self, urls):
    """Process multiple job application URLs"""
    results = []
    open_tabs = []  # Track browsers with failed submissions
    
    for i, url in enumerate(urls, 1):
      print(f"\n=== Processing URL {i}/{len(urls)}: {url} ===")
      try:
        # Initialize new ApplicationFiller instance for each URL
        filler = ApplicationFiller(
          link=url,
          headless=self.headless,
          slow_mode=self.slow_mode,
          model=self.model,
          resume_path=self.resume_path,
          reference_doc_path=self.reference_doc_path
        )
        result = filler.submit(multi_url_mode=True)
        
        # Check if submission failed (returns URL string instead of True)
        if isinstance(result, str):
          print(f"⚠️ Submission failed for {url} - keeping tab open for manual completion")
          results.append({"url": url, "status": "partial", "result": result, "driver": filler.driver})
          open_tabs.append({"url": url, "driver": filler.driver, "filler": filler})
          # Don't close the driver - keep it open for manual completion
        else:
          print(f"✓ Successfully processed {url}")
          results.append({"url": url, "status": "success", "result": result})
          filler.close_driver()
          
      except Exception as e:
        print(f"✗ Failed to process {url}: {str(e)}")
        results.append({"url": url, "status": "error", "error": str(e)})
        # Close driver on hard errors
        try:
          filler.close_driver()
        except:
          pass
      
      # Add delay between applications to avoid rate limiting
      if i < len(urls):
        print(f"Waiting 10 seconds before next application...")
        time.sleep(10)
    
    # Report summary
    print(f"\n📊 BATCH PROCESSING COMPLETE:")
    success_count = sum(1 for r in results if r['status'] == 'success')
    partial_count = sum(1 for r in results if r['status'] == 'partial')
    error_count = sum(1 for r in results if r['status'] == 'error')
    
    print(f"✅ Fully completed: {success_count}/{len(results)}")
    print(f"⚠️ Partially completed (manual finish needed): {partial_count}/{len(results)}")
    print(f"❌ Failed: {error_count}/{len(results)}")
    
    if open_tabs:
      print(f"\n🌐 {len(open_tabs)} browser tabs remain open for manual completion:")
      for tab in open_tabs:
        print(f"   📋 {tab['url']}")
      
      print(f"\n⏰ All tabs will remain open. Complete them manually and close when done.")
      print(f"   Press Ctrl+C to close all remaining tabs and exit.")
      
      # Keep script running so tabs stay open
      try:
        while open_tabs:
          time.sleep(10)
          # Check if any tabs were manually closed
          remaining_tabs = []
          for tab in open_tabs:
            try:
              # Test if driver is still active
              tab['driver'].current_url
              remaining_tabs.append(tab)
            except:
              # Tab was closed
              print(f"✓ Tab closed for {tab['url']}")
          open_tabs = remaining_tabs
          
      except KeyboardInterrupt:
        print(f"\n👋 Closing all remaining tabs...")
        for tab in open_tabs:
          try:
            tab['filler'].close_driver()
          except:
            pass
    
    return results

  def create_account_if_needed(self):
    """Create account if the site requires one and we're not logged in"""
    try:
      # Get site config based on current URL
      site_config = self.get_site_config()
      
      if not site_config.get("requires_account", False):
        return True
      
      # Check if already logged in
      if self.is_logged_in():
        print("Already logged in")
        return True
      
      print("Account required. Attempting to create/login...")
      
      # Try to login first
      if self.attempt_login(site_config):
        return True
      
      # If login fails, try account creation
      return self.attempt_account_creation(site_config)
    except Exception as e:
      print(f"Error in account creation: {e}")
      return True  # Continue anyway

  def get_site_config(self):
    """Get configuration for current site"""
    try:
      current_url = self.driver.current_url.lower()
      for domain, config in self.website_configs.items():
        if domain in current_url:
          return config
      return {}  # Return empty config if no match
    except:
      return {}

  def is_logged_in(self):
    """Check if user is already logged in"""
    try:
      # Look for common login indicators
      login_indicators = [
        "logout", "sign out", "profile", "dashboard", 
        "my account", "welcome back", "apply now"
      ]
      
      for indicator in login_indicators:
        elements = self.driver.find_elements(By.XPATH, f"//*[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{indicator}')]")
        if elements:
          print(f"Found login indicator: {indicator}")
          return True
      
      return False
    except Exception as e:
      print(f"Error checking login status: {e}")
      return False

  def attempt_login(self, site_config):
    """Attempt to login with existing credentials"""
    try:
      login_url = site_config.get("login_url")
      if login_url:
        current_url = self.driver.current_url
        # Navigate to login if not already there
        if "login" not in current_url.lower():
          base_url = self.driver.current_url.split('/')[0] + '//' + self.driver.current_url.split('/')[2]
          self.driver.get(base_url + login_url)
          time.sleep(3)
      
      # Look for login form
      email_field = site_config.get("email_field", "email")
      password_field = site_config.get("password_field", "password")
      
      # Try various selectors for email field
      email_selectors = [
        f"input[name='{email_field}']",
        f"input[id='{email_field}']",
        "input[type='email']",
        "input[placeholder*='email']"
      ]
      
      email_element = None
      for selector in email_selectors:
        try:
          email_element = self.driver.find_element(By.CSS_SELECTOR, selector)
          break
        except:
          continue
      
      if not email_element:
        print("Could not find email field for login")
        return False
      
      # Try various selectors for password field  
      password_selectors = [
        f"input[name='{password_field}']",
        f"input[id='{password_field}']",
        "input[type='password']",
        "input[placeholder*='password']"
      ]
      
      password_element = None
      for selector in password_selectors:
        try:
          password_element = self.driver.find_element(By.CSS_SELECTOR, selector)
          break
        except:
          continue
      
      if not password_element:
        print("Could not find password field for login")
        return False
      
      # Fill login form
      email_element.clear()
      email_element.send_keys(self.responses["email"])
      time.sleep(1)
      
      password_element.clear()
      password_element.send_keys(self.responses["account_password"])
      time.sleep(1)
      
      # Submit form
      submit_selectors = [
        "button[type='submit']",
        "input[type='submit']",
        "button:contains('Login')",
        "button:contains('Sign in')",
        "button:contains('Log in')"
      ]
      
      for selector in submit_selectors:
        try:
          submit_button = self.driver.find_element(By.CSS_SELECTOR, selector)
          submit_button.click()
          break
        except:
          continue
      
      # Wait for login to complete
      time.sleep(5)
      
      # Check if login was successful
      return self.is_logged_in()
      
    except Exception as e:
      print(f"Error during login attempt: {e}")
      return False

  def attempt_account_creation(self, site_config):
    """Attempt to create a new account"""
    try:
      account_url = site_config.get("account_creation_url")
      if account_url:
        base_url = self.driver.current_url.split('/')[0] + '//' + self.driver.current_url.split('/')[2]
        self.driver.get(base_url + account_url)
        time.sleep(3)
      
      # Look for registration/signup links
      signup_links = [
        "a:contains('Sign up')",
        "a:contains('Register')",
        "a:contains('Create account')",
        "button:contains('Sign up')",
        "button:contains('Register')"
      ]
      
      for link_selector in signup_links:
        try:
          link = self.driver.find_element(By.CSS_SELECTOR, link_selector)
          link.click()
          time.sleep(3)
          break
        except:
          continue
      
      # Fill registration form
      self.fill_registration_form()
      
      # Wait for account creation
      time.sleep(10)
      
      return self.is_logged_in()
      
    except Exception as e:
      print(f"Error during account creation: {e}")
      return False

  def fill_registration_form(self):
    """Fill out registration form with user details"""
    try:
      # Common registration fields
      field_mappings = {
        "email": self.responses["email"],
        "password": self.responses["account_password"],
        "confirm_password": self.responses["account_password"],
        "first_name": self.responses["first_name"],
        "last_name": self.responses["last_name"],
        "phone": self.responses["phone"],
        "full_name": f"{self.responses['first_name']} {self.responses['last_name']}"
      }
      
      for field_name, value in field_mappings.items():
        selectors = [
          f"input[name='{field_name}']",
          f"input[id='{field_name}']",
          f"input[placeholder*='{field_name.replace('_', ' ')}']"
        ]
        
        for selector in selectors:
          try:
            element = self.driver.find_element(By.CSS_SELECTOR, selector)
            element.clear()
            element.send_keys(value)
            time.sleep(0.5)
            break
          except:
            continue
      
      # Handle checkboxes (terms, privacy policy)
      checkbox_selectors = [
        "input[type='checkbox']",
        "input[name*='terms']",
        "input[name*='privacy']",
        "input[name*='agree']"
      ]
      
      for selector in checkbox_selectors:
        try:
          checkboxes = self.driver.find_elements(By.CSS_SELECTOR, selector)
          for checkbox in checkboxes:
            if not checkbox.is_selected():
              checkbox.click()
              time.sleep(0.5)
        except:
          continue
      
      # Submit registration form
      submit_selectors = [
        "button[type='submit']",
        "input[type='submit']",
        "button:contains('Create account')",
        "button:contains('Register')",
        "button:contains('Sign up')"
      ]
      
      for selector in submit_selectors:
        try:
          submit_button = self.driver.find_element(By.CSS_SELECTOR, selector)
          submit_button.click()
          break
        except:
          continue
      
    except Exception as e:
      print(f"Error filling registration form: {e}")

  def handle_education_fields(self):
    """Handle education-specific form fields"""
    try:
      education_fields = {
        "university": "Imperial College London",
        "degree": "Master of Engineering (MEng)",
        "field_of_study": "Computer Science",
        "graduation_year": "2024",
        "gpa": "First Class Honours",
        "education_level": "Master's degree"
      }
      
      for field_name, value in education_fields.items():
        selectors = [
          f"input[name*='{field_name}']",
          f"select[name*='{field_name}']",
          f"input[id*='{field_name}']",
          f"select[id*='{field_name}']"
        ]
        
        for selector in selectors:
          try:
            elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
            for element in elements:
              if element.tag_name == "select":
                self.handle_select_element(element, value)
              else:
                element.clear()
                element.send_keys(value)
              time.sleep(0.5)
              break
          except:
            continue
            
    except Exception as e:
      print(f"Error handling education fields: {e}")

  def handle_experience_fields(self):
    """Handle work experience-specific form fields"""
    try:
      # Handle multiple work experiences
      experiences = [
        {
          "company": "Trajex",
          "position": "Machine Learning Developer",
          "start_date": "June 2024",
          "end_date": "August 2024",
          "description": "Developed ML models for trading algorithms"
        },
        {
          "company": "Altus Reach", 
          "position": "Machine Learning Engineer",
          "start_date": "September 2023",
          "end_date": "November 2023", 
          "description": "Built predictive models for financial markets"
        }
      ]
      
      # Look for experience sections
      for i, exp in enumerate(experiences):
        section_prefix = f"experience_{i}" if i > 0 else "experience"
        
        for field_name, value in exp.items():
          selectors = [
            f"input[name*='{section_prefix}_{field_name}']",
            f"input[name*='{field_name}']",
            f"textarea[name*='{field_name}']",
            f"select[name*='{field_name}']"
          ]
          
          for selector in selectors:
            try:
              elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
              for element in elements[:1]:  # Only fill first match
                if element.tag_name == "select":
                  self.handle_select_element(element, value)
                elif element.tag_name == "textarea":
                  element.clear()
                  element.send_keys(value)
                else:
                  element.clear()
                  element.send_keys(value)
                time.sleep(0.5)
                break
            except:
              continue
              
    except Exception as e:
      print(f"Error handling experience fields: {e}")

  def submit(self, multi_url_mode=False):
    """Submit the form with improved detection and keep page open on failure"""
    try:
      # First, initialize driver if needed
      if not self.driver:
        if not self.initialize_driver():
          return self.link
      
      if not self.driver:
        print("❌ Failed to initialize driver")
        return self.link
      
      print(f"🌐 Navigating to: {self.link}")
      self.driver.get(self.link)
      time.sleep(3)
      
      # Step 1: Handle cookie consent banners FIRST
      self.handle_cookies()
      
      # Step 2: Handle any redirects
      self.handle_redirects()
      
      # Step 3: Extract job information BEFORE clicking any buttons
      job_description = self.extract_job_information_early()
      
      # Step 4: Analyze job and select appropriate resume using the extracted info
      selected_resume = self.analyze_job_and_select_resume(job_description)
      self.resume_path = selected_resume
      print(f"📄 Using resume: {selected_resume}")
      
      # Step 5: Handle account creation if needed
      site_config = self.get_site_config()
      if site_config.get("requires_account", False):
        self.create_account_if_needed()
      
      # Step 6: Handle Apply button if needed (this might navigate to a new page)  
      self.handle_apply_button()
      
      # Step 7: Handle cookies again in case Apply button led to new page
      self.handle_cookies()
      
      # Step 8: Fill the form
      if self.fill_form():
        # After filling, attempt to detect and refill any missing/invalid fields
        try:
          print("🔁 Running post-fill validation and refill pass...")
          self.report_and_refill_missing_fields()
          time.sleep(0.6)
          # Second quick pass
          self.report_and_refill_missing_fields()
        except Exception as e:
          print(f"⚠️ Post-fill refill pass encountered an error: {e}")
        # Find and click submit button
        from selenium.webdriver.common.by import By
        
        if not self.driver:
          print("❌ Driver not initialized")
          if multi_url_mode:
            return self.link
          else:
            self.keep_browser_open("Driver not initialized", multi_url_mode)
            return self.link
        
        print("🔍 Looking for submit button...")
        
        # Store original URL to compare later
        original_url = self.driver.current_url
        
        # Try different submit button selectors
        submit_selectors = [
          "button[type='submit']",
          "input[type='submit']",
          "button:contains('Submit')",
          "button:contains('Apply')",
          "[data-testid*='submit']",
          ".submit-button",
          "#submit",
          "input[value*='Submit']",
          "input[value*='Apply']"
        ]
        
        submit_button = None
        for selector in submit_selectors:
          try:
            submit_button = self.driver.find_element(By.CSS_SELECTOR, selector)
            if submit_button.is_displayed() and submit_button.is_enabled():
              break
          except:
            continue
        
        if submit_button and submit_button.is_displayed() and submit_button.is_enabled():
          if not self.headless and self.driver:
            # Highlight submit button
            self.driver.execute_script("arguments[0].style.border='3px solid blue'", submit_button)
            print("🎯 Found submit button - clicking in 3 seconds...")
            time.sleep(3)
          
          # Try multiple methods to click the submit button
          click_success = False
          methods_tried = []
          
          # Method 1: Regular click
          try:
            submit_button.click()
            print("✅ Submit button clicked successfully")
            click_success = True
          except Exception as e1:
            methods_tried.append(f"regular_click: {str(e1)[:50]}")
            
            # Method 2: JavaScript click
            try:
              if self.driver:
                self.driver.execute_script("arguments[0].click();", submit_button)
                print("✅ Form submitted successfully with JavaScript click")
                click_success = True
            except Exception as e2:
              methods_tried.append(f"js_click: {str(e2)[:50]}")
              
              # Method 3: Scroll into view and click
              try:
                if self.driver:
                  self.driver.execute_script("arguments[0].scrollIntoView(true);", submit_button)
                  time.sleep(1)
                  submit_button.click()
                  print("✅ Form submitted successfully after scrolling")
                  click_success = True
              except Exception as e3:
                methods_tried.append(f"scroll_click: {str(e3)[:50]}")
                
                # Method 4: Move to element and click
                try:
                  from selenium.webdriver.common.action_chains import ActionChains
                  actions = ActionChains(self.driver)
                  actions.move_to_element(submit_button).click().perform()
                  print("✅ Form submitted successfully with ActionChains")
                  click_success = True
                except Exception as e4:
                  methods_tried.append(f"action_click: {str(e4)[:50]}")
          
          if click_success:
            # Wait and check if submission was actually successful
            print("⏳ Waiting to verify submission...")
            time.sleep(3)
            
            # Check for submission success indicators
            submission_success = self.check_submission_success(original_url)
            
            if submission_success:
              print("🎉 Application submitted successfully!")
              if not self.headless:
                time.sleep(5)  # Show success for a bit
              return True
            else:
              print("⚠️ Submission may have failed - keeping browser open for manual completion")
              self.keep_browser_open("Submission verification failed", multi_url_mode)
              return self.link
          else:
            print(f"❌ All submit methods failed: {', '.join(methods_tried)}")
            self.keep_browser_open(f"Submit failed: {', '.join(methods_tried)}", multi_url_mode)
            return self.link
          
        else:
          print("⚠️ Submit button not found or not clickable")
          self.keep_browser_open("Submit button not found", multi_url_mode)
          return self.link
      else:
        print("❌ Form filling failed")
        self.keep_browser_open("Form filling failed", multi_url_mode)
        return self.link
      
    except Exception as e:
      print(f"❌ Error during submission: {e}")
      self.keep_browser_open(f"Submission error: {e}", multi_url_mode)
      return self.link

  def check_submission_success(self, original_url):
    """Check if form submission was successful"""
    try:
      from selenium.webdriver.common.by import By
      
      if not self.driver:
        return False
      
      current_url = self.driver.current_url
      
      # Check if URL changed (usually indicates successful submission)
      if current_url != original_url:
        print(f"✅ URL changed from {original_url} to {current_url}")
        return True
      
      # Check for success messages
      success_indicators = [
        "thank you",
        "application submitted",
        "successfully submitted", 
        "application received",
        "thank you for applying",
        "application complete",
        "success"
      ]
      
      # Check page text for success indicators
      try:
        page_text = self.driver.find_element(By.TAG_NAME, "body").text.lower()
        for indicator in success_indicators:
          if indicator in page_text:
            print(f"✅ Found success indicator: '{indicator}'")
            return True
      except:
        pass
      
      # Check if form fields disappeared (might indicate submission)
      try:
        form_fields = self.driver.find_elements(By.CSS_SELECTOR, "input[type='text'], input[type='email'], textarea")
        if len(form_fields) == 0:
          print("✅ Form fields disappeared - likely successful submission")
          return True
      except:
        pass
      
      # Check for error messages that would indicate failed submission
      error_indicators = [
        "required",
        "please fill",
        "this field is required",
        "error",
        "invalid"
      ]
      
      try:
        page_text = self.driver.find_element(By.TAG_NAME, "body").text.lower()
        for indicator in error_indicators:
          if indicator in page_text:
            print(f"⚠️ Found error indicator: '{indicator}' - submission likely failed")
            return False
      except:
        pass
      
      print("⚠️ Could not determine submission status - assuming failed")
      return False
      
    except Exception as e:
      print(f"❌ Error checking submission success: {e}")
      return False

  def keep_browser_open(self, reason, multi_url_mode=False):
    """Keep browser open for manual completion"""
    try:
      if not self.headless and self.driver:
        print(f"\n🔄 KEEPING BROWSER OPEN: {reason}")
        print("📋 Please complete the application manually at:")
        print(f"   {self.link}")
        
        if multi_url_mode:
          print("\n⏰ Browser tab will stay open. Continuing to next URL...")
          return  # Don't block in multi-URL mode
        else:
          print("\n⏰ Browser will stay open. Close when done or press Ctrl+C to exit.")
          
          # Keep the browser open indefinitely only in single-URL mode
          while True:
            try:
              time.sleep(10)
              # Check if browser is still open
              if not self.driver.current_url:
                break
            except:
              break
            
    except KeyboardInterrupt:
      print("\n👋 Exiting as requested")
    except Exception as e:
      print(f"⚠️ Error keeping browser open: {e}")

  def cleanup(self):
    """Clean up the web driver"""
    if self.driver and not self.headless:
      print("⏸️  Keeping browser open for 10 seconds to see results...")
      time.sleep(10)
    if self.driver:
      self.driver.quit()
  
  def close_driver(self):
    """Clean up the web driver"""
    if self.driver:
      self.driver.quit()
      self.driver = None


if __name__ == "__main__":
  # Example usage with enhanced features - LIVE MODE by default
  
  # Multi-URL support examples:
  target_urls = [
    "https://job-boards.greenhouse.io/point72/jobs/8018862002?gh_jid=8018862002&gh_src=97fa02a42us&jobCode=CSS-0013383&location=New+York",
    "https://helsing.ai/jobs/4489089101",
    "https://jobs.temasek.com.sg/job/Singapore-Vice-President%2C-Investments-117434/",
    "https://www.squarepoint-capital.com/careers/quant-developer-python/",
    "https://verition.com/careers/"
  ]
  
  # Multi-website support examples:
  websites = [
    {
      "name": "Point72 (Greenhouse)",
      "url": "https://job-boards.greenhouse.io/point72/jobs/8018862002?gh_jid=8018862002&gh_src=97fa02a42us&jobCode=CSS-0013383&location=New+York",
      "company": "Point72",
      "title": "Software Engineer"
    },
    {
      "name": "Helsing AI",
      "url": "https://helsing.ai/jobs/4489089101",
      "company": "Helsing",
      "title": "AI Research Intern"
    },
    {
      "name": "Revolut",
      "url": "https://www.revolut.com/careers/position/marketing-manager-f82b7f48-1185-4be1-b004-5131fe0ca519/",
      "company": "Revolut", 
      "title": "Marketing Manager"
    },
    {
      "name": "Da Vinci Trading",
      "url": "https://davincitrading.com/job/quant-trading-intern/",
      "company": "Da Vinci",
      "title": "Quant Trading Intern"
    },
    {
      "name": "Temasek",
      "url": "https://jobs.temasek.com.sg/job/Singapore-Vice-President%2C-Investments-117434/",
      "company": "Temasek",
      "title": "Vice President, Investments"
    }
  ]
  
  # Select processing mode
  import sys
  if len(sys.argv) > 1 and sys.argv[1] == "--multi":
    # Multi-URL processing mode
    print("🚀 Starting Multi-URL Application Processing...")
    filler = ApplicationFiller(
      link="",  # Will be set for each URL
      headless=False,
      slow_mode=True,
      model="llama3.2",
      resume_path="AdiPrabs_SWE.docx"
    )
    
    results = filler.run_multiple_applications(target_urls)
    
    print(f"\n📊 RESULTS SUMMARY:")
    success_count = sum(1 for r in results if r['status'] == 'success')
    print(f"✅ Successful: {success_count}/{len(results)}")
    print(f"❌ Failed: {len(results) - success_count}/{len(results)}")
    
    for result in results:
      status_icon = "✅" if result['status'] == 'success' else "❌"
      print(f"{status_icon} {result['url']}")
      if result['status'] == 'error':
        print(f"   Error: {result['error']}")
    
  else:
    # Single website mode
    # Select which website to use (change index to test different sites)
    selected_site = websites[0]  # Default to Point72
    
    print(f"🚀 Starting Enhanced ApplicationFiller for {selected_site['name']}...")
    print(f"🎯 Target: {selected_site['company']} - {selected_site['title']}")
    
    # Create an instance with enhanced features
    filler = ApplicationFiller(
      link=selected_site["url"],
      headless=False,          # LIVE MODE - show browser by default
      slow_mode=True,          # Slow mode for visual effects
      model="llama3.2",        # Ollama model to use
      resume_path="AdiPrabs_SWE.docx",  # Default resume (will be auto-selected based on job type)
      reference_doc_path=None,  # Path to reference document (optional)
      company_name=selected_site["company"],
      job_title=selected_site["title"]
    )
    
    # Try to fill and submit the form
    result = filler.submit(multi_url_mode=False)  # Single URL mode
    
    if isinstance(result, str):
      print(f"📋 Please complete the application manually at: {result}")
    else:
      print("🎉 Application submitted successfully!")
    
    # Clean up
    filler.close_driver()
  
  print("\n" + "="*60)
  print("✨ ENHANCED APPLICATION FILLER FEATURES:")
  print("🔹 Multi-website support (Greenhouse, Helsing, Revolut, Da Vinci, Temasek)")
  print("🔹 Smart job analysis for automatic resume selection")  
  print("🔹 Enhanced dropdown handling with Enter key support")
  print("🔹 🆕 Cookie consent handling (auto-accept)")
  print("🔹 🆕 Early job information extraction")
  print("🔹 🆕 Redirect detection and handling")
  print("🔹 🆕 Website-specific field mappings")
  print("🔹 🆕 Improved process order for better reliability")
  print("🔹 🆕 Account creation for sites that require it")
  print("🔹 🆕 Multi-URL batch processing (use --multi flag)")
  print("🔹 Apply button detection and handling")
  print("🔹 Browser stays open on errors for manual completion")
  print("🔹 Ollama LLM integration for intelligent responses")
  print("🔹 Visual feedback in LIVE MODE")
  print("="*60)
  print("\n🔧 PROCESS ORDER (NEW & IMPROVED):")
  print("1. 🌐 Navigate to URL")
  print("2. 🍪 Handle cookie consent banners")
  print("3. 🔄 Detect and handle redirects")
  print("4. 📋 Extract job information EARLY (before Apply button)")
  print("5. 🤖 Analyze job for smart resume selection")
  print("6. 🔐 Create account if required")
  print("7. 🎯 Handle Apply button (if needed)")
  print("8. 🍪 Handle cookies again (if new page loaded)")
  print("9. 📝 Fill form with enhanced field detection")
  print("10. ✅ Submit application")
  print("="*60)
  print("\n🚀 USAGE:")
  print("Single URL: python3 ApplicationFiller.py")
  print("Multi URLs: python3 ApplicationFiller.py --multi")
  print("="*60)
