#!/usr/bin/env python3
"""
Fix Resume Customization
Update resume customization to use actual company names and job titles from job descriptions.
"""

import json
import os
from docx import Document
import re

def extract_company_and_job_from_url(url, company_name, job_title):
    """Extract actual company name and job title from URL and job description"""
    
    # Parse company name from URL or use provided company name
    actual_company = company_name
    
    # Extract company from common URL patterns
    if 'greenhouse.io' in url:
        # Extract company from greenhouse URLs
        match = re.search(r'greenhouse\.io/([^/]+)', url)
        if match:
            actual_company = match.group(1).replace('-', ' ').title()
    
    elif 'smartrecruiters.com' in url:
        # Extract company from SmartRecruiters URLs
        match = re.search(r'smartrecruiters\.com/([^/]+)', url)
        if match:
            actual_company = match.group(1).replace('-', ' ').title()
    
    elif 'workdayjobs.com' in url:
        # Extract company from Workday URLs
        match = re.search(r'workdayjobs\.com/[^/]+/([^/]+)', url)
        if match:
            actual_company = match.group(1).replace('_', ' ').title()
    
    elif 'lever.co' in url:
        # Extract company from Lever URLs
        match = re.search(r'lever\.co/([^/?]+)', url)
        if match:
            actual_company = match.group(1).replace('-', ' ').title()
    
    elif 'jobs.' in url:
        # Extract company from jobs.* URLs
        match = re.search(r'jobs\.([^/]+)', url)
        if match:
            actual_company = match.group(1).replace('.com', '').replace('.co.uk', '').title()
    
    # Clean up company name
    actual_company = re.sub(r'[^\w\s]', '', actual_company).strip()
    
    # Extract job title from job_title or URL
    actual_job_title = job_title
    
    # Clean up job title
    if actual_job_title:
        # Remove common suffixes
        actual_job_title = re.sub(r'\s*\([^)]*\)\s*', '', actual_job_title)
        actual_job_title = re.sub(r'\s*-\s*[^-]*$', '', actual_job_title)
        actual_job_title = actual_job_title.strip()
    
    return actual_company, actual_job_title

def customize_resume_with_actual_data(resume_path, company_name, job_title, url):
    """Customize resume with actual company name and job title"""
    
    if not os.path.exists(resume_path):
        print(f"❌ Resume file not found: {resume_path}")
        return resume_path
    
    # Extract actual company and job information
    actual_company, actual_job_title = extract_company_and_job_from_url(url, company_name, job_title)
    
    print(f"🏢 Company: {actual_company}")
    print(f"💼 Job Title: {actual_job_title}")
    
    if not actual_company and not actual_job_title:
        print("⚠️ No company name or job title available for customization")
        return resume_path
    
    # Create customized resumes folder
    customized_folder = "customized_resumes"
    if not os.path.exists(customized_folder):
        os.makedirs(customized_folder)
        print(f"📁 Created folder: {customized_folder}")
    
    # Create customized filename in the folder
    base_name = os.path.splitext(os.path.basename(resume_path))[0]
    extension = os.path.splitext(resume_path)[1]
    
    # Create safe filename
    safe_company_name = actual_company.replace(' ', '').replace('/', '_').replace('\\', '_')[:30]
    safe_job_title = actual_job_title.replace(' ', '').replace('/', '_').replace('\\', '_')[:30] if actual_job_title else ''
    
    if safe_job_title:
        customized_filename = f"{base_name}_{safe_company_name}_{safe_job_title}{extension}"
    else:
        customized_filename = f"{base_name}_{safe_company_name}{extension}"
    
    customized_path = os.path.join(customized_folder, customized_filename)
    
    # For DOCX files, we need python-docx library
    if resume_path.endswith('.docx'):
        try:
            # Load the document
            doc = Document(resume_path)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                if '<<Company>>' in paragraph.text:
                    paragraph.text = paragraph.text.replace('<<Company>>', actual_company)
                if '<<Program>>' in paragraph.text:
                    paragraph.text = paragraph.text.replace('<<Program>>', actual_job_title)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '<<Company>>' in paragraph.text:
                                paragraph.text = paragraph.text.replace('<<Company>>', actual_company)
                            if '<<Program>>' in paragraph.text:
                                paragraph.text = paragraph.text.replace('<<Program>>', actual_job_title)
            
            # Save the customized document
            doc.save(customized_path)
            print(f"✅ Customized resume saved: {customized_path}")
            return customized_path
            
        except Exception as e:
            print(f"❌ Error customizing DOCX resume: {e}")
            return resume_path
    
    # For other file types, just copy the file
    else:
        try:
            import shutil
            shutil.copy2(resume_path, customized_path)
            print(f"✅ Resume copied: {customized_path}")
            return customized_path
        except Exception as e:
            print(f"❌ Error copying resume: {e}")
            return resume_path

def test_resume_customization():
    """Test the resume customization with sample data"""
    print("🧪 Testing Resume Customization")
    print("=" * 50)
    
    # Test cases with real URLs and job data
    test_cases = [
        {
            "resume_path": "AdiPrabs_Quant.docx",
            "company": "Point72",
            "job_title": "2026 Academy Investment",
            "url": "https://job-boards.greenhouse.io/point72/jobs/7885446002"
        },
        {
            "resume_path": "AdiPrabs_SWE.docx",
            "company": "Jane Street",
            "job_title": "Quantitative Trader/Researcher Internship",
            "url": "https://www.janestreet.com/join-jane-street/open-roles/?type=internship"
        },
        {
            "resume_path": "AdiPrabs_Cons.docx",
            "company": "Apple",
            "job_title": "Finance Development Program",
            "url": "https://jobs.apple.com/en-us/details/200608579/uk-finance-development-program"
        }
    ]
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\n📋 Test Case {i}:")
        print(f"   Resume: {test_case['resume_path']}")
        print(f"   Company: {test_case['company']}")
        print(f"   Job: {test_case['job_title']}")
        print(f"   URL: {test_case['url']}")
        
        result = customize_resume_with_actual_data(
            test_case['resume_path'],
            test_case['company'],
            test_case['job_title'],
            test_case['url']
        )
        
        print(f"   Result: {result}")

def main():
    """Main function"""
    print("🔧 Fixing Resume Customization")
    print("=" * 50)
    
    # Test the customization
    test_resume_customization()
    
    print(f"\n✅ Resume customization fixes completed!")

if __name__ == "__main__":
    main()

